#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
M15 - Redacción del informe en Word
===================================
Entorno: venv del proyecto.

RESUELVE LO MECÁNICO Y NO ESCRIBE JUICIO. La plantilla del consultor lleva sus
instrucciones marcadas en verde, de tres tipos:

    Colocar Figura: <archivo>.png    inserta la figura que produjo la cadena
    Completar la Tabla N-M           llena la tabla con los datos del estudio
    Analizar Ilustración N-M         redacción, que este módulo NO toca

Las dos primeras son mecánicas: el archivo existe o no, el dato está o no. La
tercera exige mirar el resultado y decir qué significa, y eso no se programa. El
módulo las deja intactas, en verde, para que la pasada de redacción las
encuentre.

SE EDITA EL DOCUMENTO, NO SE RECONSTRUYE. Es la diferencia entre conservar las
92 referencias cruzadas, las 152 leyendas numeradas por campo, los 214
hipervínculos y los 314 marcadores, o perderlos todos. Verificado: tras editar,
los recuentos de REF, SEQ, STYLEREF, PAGEREF, TOC, marcadores e hipervínculos
son idénticos.

LOS ÍNDICES QUEDAN DESACTUALIZADOS a propósito. Son campos de Word y conservan
su texto en caché hasta que se actualizan: al abrir el documento hay que
responder que sí, o pulsar Ctrl+E y F9.

Productos:
    <salida>/informe_hidrologico.docx
    data/02_procesado/M15_informe.json

Uso:
    python src/M15_informe.py
    python src/M15_informe.py --plantilla otra_base.docx

Códigos de salida:
    0  informe escrito sin hallazgos bloqueantes
    1  hay hallazgos bloqueantes
    3  no se pudo leer la configuración o la plantilla
"""

from __future__ import annotations

import argparse
import csv
import json
import re
import sys
import time
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Sequence

_DIRECTORIO_SRC = Path(__file__).resolve().parent
if str(_DIRECTORIO_SRC) not in sys.path:
    sys.path.insert(0, str(_DIRECTORIO_SRC))

import docx_plantilla as plantilla_docx  # noqa: E402
from comun import esquema, registro, rutas  # noqa: E402
from comun.config import Config, cargar, leer_yaml  # noqa: E402
from comun.errores import (  # noqa: E402
    ErrorConfiguracion,
    ErrorFormato,
    ErrorHidrologia,
    ErrorRutas,
)
from comun.esquema import ADVERTENCIA, BLOQUEANTE, INFORMATIVO, Hallazgo  # noqa: E402

MODULO = "M15"
DESCRIPCION = "Redacción del informe en Word"

SALIDA_CORRECTA = 0
SALIDA_BLOQUEANTE = 1
SALIDA_ERROR = 3

# Las tres instrucciones que la plantilla marca en verde. El módulo resuelve las
# dos primeras y respeta la tercera.
PATRON_FIGURA = re.compile(r"^\s*Colocar\s+Figura\s*:\s*(.+?)\s*$", re.I)
PATRON_TABLA = re.compile(r"^\s*Completar\s+la\s+Tabla\s+([0-9]+[-‐-―]?[0-9]*)",
                          re.I)
PATRON_ANALISIS = re.compile(r"^\s*Analizar\b", re.I)

# Las leyendas con que la plantilla nombra sus figuras y sus tablas. Dicen
# "Ilustracion -." mientras Word no actualice los campos SEQ, de modo que lo
# que identifica no es el numero sino el texto que sigue.
PATRON_LEYENDA = re.compile(r"^\s*(Ilustraci[oó]n|Gr[aá]fico|Figura|Tabla)\b",
                            re.I)

# Un nombre de archivo dentro de la instrucción, con su carpeta si la lleva.
# \w ES UNICODE en Python y admite tilde y ene, [A-Za-z0-9] no. Con la clase
# ASCII, un nombre como "M14_comparacion_cambio_climatico.png" escrito con
# tildes se leia TRUNCADO desde la ultima letra acentuada, "tico.png", y el
# modulo buscaba un archivo con ese nombre: no lo encontraba y reportaba una
# figura ausente en lugar de un nombre que no puede leer.
PATRON_ARCHIVO = re.compile(r"([\w\-\.]+\.(?:png|jpg|jpeg|svg))", re.I)


@dataclass
class ResultadoM15:
    """Lo que el módulo resolvió y lo que dejó pendiente."""

    plantilla: str = ""
    documento: str = ""
    figuras_puestas: list[str] = field(default_factory=list)
    figuras_ausentes: list[dict[str, Any]] = field(default_factory=list)
    figuras_corregidas: list[dict[str, Any]] = field(default_factory=list)
    correcciones_ambiguas: list[str] = field(default_factory=list)
    correcciones_sin_uso: list[str] = field(default_factory=list)
    tablas_llenadas: list[dict[str, Any]] = field(default_factory=list)
    tablas_sin_fuente: list[str] = field(default_factory=list)
    analisis_pendientes: int = 0
    productos: list[str] = field(default_factory=list)
    hallazgos: list[Hallazgo] = field(default_factory=list)


# =============================================================================
# Lectura de la plantilla
# =============================================================================
def marcar_campos_para_actualizar(documento) -> int:
    """
    Marca los campos de Word como sucios, para que se recalculen al abrir.

    POR QUE HACE FALTA. Las leyendas se numeran con campos SEQ y los indices de
    contenido, de ilustraciones y de tablas son campos TOC. Todos conservan su
    ULTIMO RESULTADO EN CACHE, que es el del informe del que se derivo la
    plantilla: por eso una leyenda dice "Tabla -." y otra "Tabla 5-10" cuando
    en el documento le corresponde otro numero. Insertar una tabla nueva no
    renumera nada por si solo.

    El atributo 'w:dirty' es la forma que el formato prevé para decir que el
    resultado guardado ya no vale. Word lo recalcula al abrir el documento; con
    eso las leyendas nuevas entran en la secuencia de su capitulo y los indices
    recogen lo que se anadio, sin que haya que pulsar Ctrl+E y F9.

    NO TODO VISOR LO HONRA. LibreOffice y las vistas previas web muestran el
    valor en cache, de modo que el numero correcto se ve al abrirlo en Word.
    Por eso el aviso de que hay que actualizar no desaparece: se rebaja a
    recordatorio.

    Devuelve cuantos campos se marcaron.
    """
    from docx.oxml.ns import qn

    marcados = 0
    for elemento in documento.element.body.iter():
        if elemento.tag == qn("w:fldChar"):
            # Solo el de apertura: el separador y el de cierre no llevan la
            # instruccion, y marcarlos no significa nada.
            if elemento.get(qn("w:fldCharType")) != "begin":
                continue
        elif elemento.tag != qn("w:fldSimple"):
            continue
        elemento.set(qn("w:dirty"), "true")
        marcados += 1
    return marcados


def clasificar(texto: str) -> tuple[str, str]:
    """
    Tipo de instrucción y su argumento.

    Devuelve ('figura'|'tabla'|'analisis'|'', argumento). Se decide por el texto
    y no por el color: el sombreado se pierde al copiar y pegar un párrafo, y
    entonces una instrucción quedaría muda sin que nada lo señalara.
    """
    coincide = PATRON_FIGURA.match(texto)
    if coincide:
        archivo = PATRON_ARCHIVO.search(coincide.group(1))
        return "figura", (archivo.group(1) if archivo else coincide.group(1))
    coincide = PATRON_TABLA.match(texto)
    if coincide:
        return "tabla", coincide.group(1)
    if PATRON_ANALISIS.match(texto):
        return "analisis", texto.strip()
    return "", ""


def buscar_figura(nombre: str, raices: Sequence[Path]) -> Path | None:
    """
    Localiza una figura por su nombre, en el directorio de gráficos y sus temas.

    SE BUSCA POR NOMBRE Y NO POR RUTA. La instrucción de la plantilla dice el
    archivo y a veces la carpeta ('de la carpeta individuales/isoyetas_fase'),
    pero el consultor no tiene por qué llevar la cuenta de en qué subcarpeta lo
    dejó cada módulo: el nombre es único y basta.
    """
    for raiz in raices:
        raiz = Path(raiz)
        if not raiz.is_dir():
            continue
        directo = raiz / nombre
        if directo.is_file():
            return directo
        for encontrado in raiz.rglob(nombre):
            if encontrado.is_file():
                return encontrado
    return None


def leer_tabla_csv(ruta: Path, delimitador: str) -> list[dict[str, str]]:
    """Filas de un CSV de la cadena, con su codificación."""
    ruta = Path(ruta)
    if not ruta.is_file():
        raise ErrorRutas(f"no se encuentra la tabla {ruta}.")
    with ruta.open(encoding="utf-8-sig", newline="") as archivo:
        return list(csv.DictReader(archivo, delimiter=delimitador))


def leer_declaracion_tablas(ruta: Path) -> dict[str, dict[str, Any]]:
    """
    Qué fuente alimenta cada tabla del informe, indexada por su LEYENDA.

    ES UNA DECLARACIÓN Y NO UNA REGLA. Adivinar la fuente por el texto de la
    leyenda funcionaría hasta el primer informe que llame a las cosas de otro
    modo.

    NO SE INDEXA POR EL NUMERO, y la razón es que el número no identifica nada.
    La plantilla compone sus leyendas con campos SEQ que no tienen resultado en
    caché, de modo que todas dicen literalmente 'Tabla -.' hasta que Word
    actualice los campos. El número que el consultor escribió en la instrucción
    es el del informe del que copió el apartado, no el de esta plantilla: en la
    plantilla de este estudio hay dos instrucciones que dicen 5-1 delante de
    tablas distintas, dos que dicen 5-6, ninguna que diga 5-12, y la que dice
    5-10 va delante de la tabla de tránsito de crecientes y no de la de rezago.
    Indexar por número llenaba una tabla con los datos de otra sin emitir
    ninguna señal.

    La leyenda, en cambio, dice qué es la tabla y sobrevive a que Word renumere.
    Se conserva 'numero' en la declaración como referencia para el consultor,
    pero no se usa para emparejar.
    """
    ruta = Path(ruta)
    if not ruta.is_file():
        return {}
    declaracion = leer_yaml(ruta) or {}
    salida: dict[str, dict[str, Any]] = {}
    for entrada in declaracion.get("tablas") or []:
        clave = _normalizar_leyenda(str(entrada.get("leyenda", "")))
        if clave:
            salida[clave] = entrada
    return salida


def leer_correcciones(ruta: Path) -> dict[str, dict[str, Any]]:
    """
    Qué instrucciones de la plantilla piden lo que no corresponde.

    LA PLANTILLA DEL CONSULTOR NO SE TOCA, y la copia saneada es derivada: se
    regenera desde ella y una corrección escrita allí se perdería sin aviso. Se
    declara aquí, se aplica al vuelo y se registra, de modo que el estudio
    puede explicar por qué una figura no es la que la instrucción nombraba.

    Se indexa por la leyenda por la misma razón que las tablas (ver
    'leer_declaracion_tablas'): el número no identifica nada.
    """
    ruta = Path(ruta)
    if not ruta.is_file():
        return {}
    declaracion = leer_yaml(ruta) or {}
    salida: dict[str, dict[str, Any]] = {}
    for entrada in declaracion.get("figuras") or []:
        clave = _normalizar_leyenda(str(entrada.get("leyenda", "")))
        if clave:
            salida[clave] = entrada
    return salida


def leer_sustituciones(ruta: Path) -> list[dict[str, Any]]:
    """
    Qué nombres el informe escribe de otra forma que sus fuentes.

    El catálogo del IDEAM entrega la entidad operadora con su razón social
    completa en las 44 estaciones del inventario, y el informe la nombra por su
    sigla. Es una convención de redacción del consultor, de modo que se declara
    junto a las correcciones y no se codifica.
    """
    ruta = Path(ruta)
    if not ruta.is_file():
        return []
    return list((leer_yaml(ruta) or {}).get("texto") or [])


def leyendas_de_instruccion(documento) -> dict[Any, str]:
    """
    La leyenda que precede a cada instrucción de figura.

    EN LAS FIGURAS LA LEYENDA VA ANTES Y EN LAS TABLAS DESPUES. No es un
    descuido de la plantilla sino su composición, y buscar en la dirección
    equivocada empareja cada figura con la leyenda de la siguiente.

    Se recorre el documento en orden, párrafos del cuerpo y de las celdas,
    porque dieciséis de las figuras se componen dentro de tablas para ponerlas
    de a dos y esos párrafos no están en 'documento.paragraphs'.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    salida: dict[Any, str] = {}
    ultima = ""

    def mirar(parrafo) -> None:
        nonlocal ultima
        texto = parrafo.text.strip()
        if not texto:
            return
        if clasificar(texto)[0] == "figura":
            salida[parrafo._element] = ultima
        elif PATRON_LEYENDA.match(texto):
            ultima = texto
        else:
            # LA LEYENDA TIENE QUE IR PEGADA A SU INSTRUCCION. Cualquier otro
            # párrafo la cancela, empezando por el 'Fuente:' que sigue a cada
            # figura. Sin esto, una instrucción sin leyenda propia heredaba la
            # de la figura anterior y una corrección declarada se aplicaba
            # también a figuras que no eran la suya.
            ultima = ""

    for hijo in documento.element.body.iterchildren():
        if hijo.tag.endswith("}p"):
            mirar(Paragraph(hijo, documento))
        elif hijo.tag.endswith("}tbl"):
            for fila in Table(hijo, documento).rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        mirar(parrafo)
    return salida


def planear_correcciones(documento, correcciones):
    """
    A qué instrucción concreta se aplica cada corrección declarada.

    UNA CORRECCION QUE EMPAREJA CON VARIAS NO SE APLICA A NINGUNA. En esta
    plantilla hay tres instrucciones bajo la leyenda 'Áreas microcuencas': la
    del mapa de áreas y dos figuras de pendiente cuya leyenda quedó mal
    copiada. Aplicar la corrección a las tres sustituía dos figuras correctas
    por el mapa de áreas, y el informe salía con la misma imagen repetida sin
    que nada lo dijera. Para desempatar se declara 'pedia', el archivo que la
    instrucción nombra hoy.

    Devuelve (plan, ambiguas, sin_uso): el plan por párrafo, las correcciones
    que emparejan con más de una instrucción y las que no emparejan con
    ninguna.
    """
    leyendas = leyendas_de_instruccion(documento)
    instrucciones: list[tuple[Any, str, str]] = []
    for elemento, leyenda in leyendas.items():
        texto = "".join(elemento.itertext())
        instrucciones.append((elemento, leyenda, clasificar(texto)[1]))

    plan: dict[Any, dict[str, Any]] = {}
    ambiguas: list[str] = []
    sin_uso: list[str] = []
    for clave, entrada in correcciones.items():
        pedia = str(entrada.get("pedia", "")).strip()
        candidatos = [e for e, ley, nom in instrucciones
                      if _normalizar_leyenda(ley) == clave
                      and (not pedia or nom == pedia)]
        etiqueta = str(entrada.get("leyenda", clave))
        if not candidatos:
            sin_uso.append(etiqueta + (f" (pedia {pedia})" if pedia else ""))
        elif len(candidatos) > 1:
            ambiguas.append(f"{etiqueta}: {len(candidatos)} instrucciones")
        else:
            plan[candidatos[0]] = entrada
    return plan, ambiguas, sin_uso


def _normalizar_leyenda(texto: str) -> str:
    """
    Deja la leyenda en lo que identifica al objeto y nada más.

    Se quita el prefijo numerado, que es justamente lo que no se puede
    comparar ('Tabla 2-1.', 'Tabla -.', 'Ilustración -.'), y se igualan tildes,
    mayúsculas y espacios: la declaración la escribe el consultor a mano y no
    tiene por qué reproducir la acentuación de la plantilla.

    SIRVE PARA LAS CUATRO PALABRAS con que la plantilla encabeza sus leyendas.
    Quitar solo 'Tabla' dejaba las de figura como 'ilustracion -. areas
    microcuencas', que no emparejaba con nada y hacía que la corrección
    declarada no se aplicara.
    """
    limpio = re.sub(r"^\s*(?:tabla|ilustraci[oó]n|gr[aá]fico|figura)"
                    r"\s*[0-9\-‐-―.\s]*", "", texto.strip(), flags=re.I)
    limpio = unicodedata.normalize("NFKD", limpio)
    limpio = "".join(c for c in limpio if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", limpio).strip().lower()


def _normalizar_numero(texto: str) -> str:
    """
    '2-1', '2‑1' y '21' son el mismo número de tabla.

    Word compone la leyenda con campos y el guion que aparece en pantalla puede
    ser un guion corto, uno largo o uno de no separación; y al extraer el texto
    de los campos el separador se pierde del todo.
    """
    return re.sub(r"[^0-9]", "", str(texto))


# =============================================================================
# Edición del documento
# =============================================================================
def _ancho_util(documento) -> Any:
    """Ancho de la caja de texto de la primera sección."""
    seccion = documento.sections[0]
    return seccion.page_width - seccion.left_margin - seccion.right_margin


def _ancho_de_celda(celda, ancho_pagina) -> int:
    """
    Ancho util de una celda, para que la figura no la desborde.

    UNA FIGURA A ANCHO DE PAGINA DENTRO DE UNA CELDA rompe la tabla: Word la
    ensancha hasta salirse del margen. Se toma el ancho declarado de la celda y
    se le descuenta el margen habitual; si la celda no lo declara, se reparte el
    ancho de pagina entre las columnas de su fila.
    """
    from docx.shared import Emu

    declarado = celda.width
    if declarado is not None and int(declarado) > 0:
        return max(int(Emu(360000)), int(declarado) - int(Emu(144000)))
    return int(ancho_pagina)


def poner_figura(parrafo, ruta_imagen: Path, ancho) -> None:
    """
    Sustituye el texto de la instrucción por la imagen, centrada.

    SE CONSERVA EL PÁRRAFO, no se crea otro. Su estilo es el que la plantilla
    reservó para la figura, y los párrafos vecinos son la leyenda y la fuente:
    insertar uno nuevo dejaría la figura fuera de esa secuencia.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Emu

    for run in list(parrafo.runs):
        run._element.getparent().remove(run._element)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parrafo.add_run().add_picture(str(ruta_imagen), width=Emu(int(ancho)))


def _fijar_texto(celda, valor: str) -> None:
    """
    Escribe en una celda conservando el formato del primer run que ya tenía.

    La plantilla trae las tablas del estudio anterior con su tipografía y su
    tamaño; el consultor pidió expresamente conservarlos. Reescribir la celda
    entera perdería ese formato, de modo que se reutiliza el primer run y se
    vacían los demás.
    """
    parrafo = celda.paragraphs[0]
    if parrafo.runs:
        parrafo.runs[0].text = valor
        for run in parrafo.runs[1:]:
            run.text = ""
    else:
        parrafo.add_run(valor)
    for sobrante in celda.paragraphs[1:]:
        sobrante._element.getparent().remove(sobrante._element)


def formatear_numero(texto: str, decimales: int) -> str:
    """
    Deja una cifra con los decimales que el informe usa, y lo demás intacto.

    SOLO SE TOCA LO QUE YA TRAE DECIMALES. Un año, un código de estación o un
    conteo son números y no llevan parte decimal: redondearlos a dos los
    convertiría en '1983.00' y en '21201230.00'. La regla se aplica a lo que
    tiene separador decimal y a nada más.

    UNA CIFRA PEQUEÑA NO SE CONVIERTE EN CERO. Si redondear deja 0.00 y el
    valor no era cero, se conserva como venía: un caudal de 0.0004 m3/s es un
    dato y '0.00' es una pérdida silenciosa.

    El separador es el punto, que es el que la plantilla usa (medido: 50 cifras
    con punto y una con coma).
    """
    limpio = texto.strip()
    if not limpio or "." not in limpio:
        return texto
    try:
        valor = float(limpio)
    except ValueError:
        return texto
    redondeado = f"{valor:.{decimales}f}"
    if float(redondeado) == 0.0 and valor != 0.0:
        return texto
    return redondeado


def aplicar_sustituciones(texto: str,
                          sustituciones: Sequence[dict[str, Any]]) -> str:
    """
    Cambia los nombres que el informe escribe de otra forma.

    El catálogo del IDEAM entrega la entidad operadora con su razón social
    completa, y el informe la nombra por su sigla. Es una convención de
    redacción del consultor y por eso se declara, no se codifica.
    """
    for regla in sustituciones or []:
        busca = str(regla.get("busca", ""))
        if busca and busca in texto:
            texto = texto.replace(busca, str(regla.get("pone", "")))
    return texto


def _valor(texto: Any, decimales: int, sustituciones) -> str:
    """Lo que se escribe en una celda: cifra redondeada y nombre sustituido."""
    return formatear_numero(
        aplicar_sustituciones(str(texto).strip(), sustituciones), decimales)


def anadir_columnas(tabla, cuantas: int) -> int:
    """
    Ensancha la tabla hasta 'cuantas' columnas, copiando la última.

    LA PLANTILLA VIENE DIMENSIONADA PARA CUATRO MICROCUENCAS y este estudio
    tiene 125. Las tablas de tiempo de concentración y de rezago son matrices de
    autor por unidad; con 125 unidades la única orientación que cabe en una
    página es la traspuesta, unidad por autor, y entonces hacen falta tantas
    columnas como fórmulas. Crecer en filas y no en columnas dejaba la mitad de
    la matriz fuera sin que nada lo dijera.

    SE COPIA LA ULTIMA CELDA DE CADA FILA y no se crea una vacía: así la columna
    nueva hereda bordes, sombreado y tipografía de la que ya estaba, igual que
    se hace con las filas. Devuelve cuántas se añadieron.
    """
    import copy

    from docx.oxml.ns import qn

    faltan = cuantas - len(tabla.columns)
    if faltan <= 0:
        return 0

    cuadricula = tabla._tbl.find(qn("w:tblGrid"))
    for _ in range(faltan):
        if cuadricula is not None and len(cuadricula):
            cuadricula.append(copy.deepcopy(cuadricula[-1]))
        for fila in tabla.rows:
            celdas = fila._tr.findall(qn("w:tc"))
            if celdas:
                fila._tr.append(copy.deepcopy(celdas[-1]))
    return faltan


def llenar_matriz(tabla, filas: Sequence[dict[str, str]],
                  matriz: dict[str, Any], encabezados: int,
                  decimales: int = 2, sustituciones=()) -> dict[str, Any]:
    """
    Llena una tabla cuyas columnas son valores de un campo, no campos.

    DOCE DE LAS TABLAS DEL INFORME SON MATRICES: año por mes, duración por
    periodo de retorno, autor por subcuenca. El CSV las trae en formato largo,
    una fila por celda, y la declaración dice qué campo da la etiqueta de fila,
    qué campo da las columnas y cuál el dato.

    'orden' enumera los valores de la columna EN EL ORDEN DE LA TABLA, y no se
    deduce del encabezado: la plantilla escribe '2.33' donde el CSV dice
    '2.33', pero escribe 'ENE' donde el CSV dice '1'. Deducirlo funcionaría
    hasta el primer informe que abrevie los meses de otro modo.

    Una celda sin dato queda VACIA y se reporta. Dejarla con lo que la
    plantilla traía pondría un valor de otro estudio en esta tabla.
    """
    import copy

    campo_fila = str(matriz["fila"])
    campo_columna = str(matriz["columna"])
    campo_valor = str(matriz["valor"])
    orden = [str(v) for v in matriz.get("orden") or []]

    for campo in (campo_fila, campo_columna, campo_valor):
        if filas and campo not in filas[0]:
            raise ErrorFormato(
                f"la fuente no tiene la columna '{campo}' que la matriz "
                f"declara; tiene {sorted(filas[0])}.")

    etiquetas: list[str] = []
    celdas: dict[tuple[str, str], str] = {}
    for fila in filas:
        clave_fila = str(fila.get(campo_fila, "")).strip()
        if clave_fila not in etiquetas:
            etiquetas.append(clave_fila)
        celdas[(clave_fila, str(fila.get(campo_columna, "")).strip())] = str(
            fila.get(campo_valor, ""))

    disponibles = len(tabla.rows) - encabezados
    if disponibles < 1:
        raise ErrorFormato(
            f"la tabla declara {encabezados} fila(s) de encabezado y solo "
            f"tiene {len(tabla.rows)}: no queda ninguna de datos.")

    if matriz.get("crecer_columnas"):
        anadir_columnas(tabla, 1 + len(orden))
        # EL ENCABEZADO SE ESCRIBE PORQUE LAS COLUMNAS SON NUEVAS. Una columna
        # anadida sale sin titulo, y una tabla con titulos en blanco no se
        # entiende. Se escribe la ULTIMA fila de encabezado, que es la que
        # nombra las columnas; las de arriba, si las hay, son el rotulo
        # combinado y no se tocan.
        titulos = [str(matriz.get("etiqueta_fila", ""))] + [
            str(v) for v in orden]
        celdas_titulo = tabla.rows[encabezados - 1].cells
        for posicion, titulo in enumerate(titulos):
            if posicion < len(celdas_titulo):
                _fijar_texto(celdas_titulo[posicion], titulo)

    modelo = tabla.rows[encabezados]._tr
    while len(tabla.rows) - encabezados < len(etiquetas):
        tabla._tbl.append(copy.deepcopy(modelo))
    while len(tabla.rows) - encabezados > len(etiquetas):
        ultimo = tabla.rows[-1]._tr
        ultimo.getparent().remove(ultimo)

    huecos = 0
    for indice, etiqueta in enumerate(etiquetas):
        fila_tabla = tabla.rows[encabezados + indice].cells
        _fijar_texto(fila_tabla[0], _valor(etiqueta, decimales, sustituciones))
        for posicion, columna in enumerate(orden, start=1):
            if posicion >= len(fila_tabla):
                break
            if (etiqueta, columna) not in celdas:
                huecos += 1
            _fijar_texto(fila_tabla[posicion],
                         _valor(celdas.get((etiqueta, columna), ""),
                                decimales, sustituciones))
    return {"filas": len(etiquetas), "columnas": 1 + len(orden),
            "sin_columna": [], "huecos": huecos}


def llenar_tabla(tabla, filas: Sequence[dict[str, str]],
                 columnas: Sequence[str], encabezados: int,
                 decimales: int = 2, sustituciones=(), titulos=()) -> dict[str, Any]:
    """
    Reemplaza los datos de la tabla conservando sus filas de encabezado.

    LAS FILAS SOBRANTES SE BORRAN Y LAS QUE FALTAN SE AÑADEN copiando la última
    de datos: así la fila nueva hereda bordes, sombreado y tipografía de la que
    ya estaba, en lugar de salir con el formato por defecto de Word.

    Devuelve el recuento, que es lo que el reporte necesita para decir si la
    tabla quedó con los datos de este estudio o con los del anterior.
    """
    import copy

    disponibles = len(tabla.rows) - encabezados
    if disponibles < 1:
        raise ErrorFormato(
            f"la tabla declara {encabezados} fila(s) de encabezado y solo tiene "
            f"{len(tabla.rows)}: no queda ninguna de datos que rellenar.")

    if titulos:
        # LA TABLA SE ENSANCHA Y LAS COLUMNAS NUEVAS SALEN SIN TITULO. Se
        # escribe la ULTIMA fila de encabezado, que es la que nombra las
        # columnas; las de arriba, si las hay, son el rotulo combinado.
        anadir_columnas(tabla, len(columnas))
        celdas_titulo = tabla.rows[encabezados - 1].cells
        for posicion, titulo in enumerate(titulos):
            if posicion < len(celdas_titulo):
                _fijar_texto(celdas_titulo[posicion], str(titulo))

    modelo = tabla.rows[encabezados]._tr
    while len(tabla.rows) - encabezados < len(filas):
        tabla._tbl.append(copy.deepcopy(modelo))
    while len(tabla.rows) - encabezados > len(filas):
        ultimo = tabla.rows[-1]._tr
        ultimo.getparent().remove(ultimo)

    sin_columna: set[str] = set()
    for indice, fila in enumerate(filas):
        celdas = tabla.rows[encabezados + indice].cells
        for posicion, columna in enumerate(columnas):
            if posicion >= len(celdas):
                break
            if columna not in fila:
                sin_columna.add(columna)
            _fijar_texto(celdas[posicion],
                         _valor(fila.get(columna, ""), decimales,
                                sustituciones))
    return {"filas": len(filas), "columnas": len(columnas),
            "sin_columna": sorted(sin_columna)}


# =============================================================================
# Ejecución
# =============================================================================
def ejecutar(
    raiz: Path | None = None,
    ruta_config: Path | None = None,
    ruta_json: Path | None = None,
    ruta_plantilla: Path | None = None,
    ruta_salida: Path | None = None,
    consola: bool = True,
) -> tuple[int, list[Hallazgo]]:
    """Resuelve figuras y tablas sobre la plantilla y escribe el informe."""
    inicio = time.perf_counter()

    base = Path(raiz).resolve() if raiz is not None else rutas.raiz_proyecto()
    configuracion: Config = cargar(ruta=ruta_config, raiz=base)
    logger = registro.configurar(
        MODULO, nivel=configuracion.obtener("ejecucion.nivel_log", "INFO"),
        raiz=base, consola=consola,
    )

    origen = (Path(ruta_plantilla) if ruta_plantilla is not None
              else rutas.resolver(configuracion.obtener("informe.plantilla_base"),
                                  rutas.raiz_codigo()))
    delimitador = configuracion.obtener("insumos_usuario.delimitador_csv", ";")
    graficos = rutas.resolver(configuracion.obtener("graficos.directorio"), base)
    individuales = rutas.resolver(
        configuracion.obtener("graficos.directorio_individuales"), base)
    declaracion = leer_declaracion_tablas(
        rutas.resolver(configuracion.obtener("informe.tablas"), base))
    ruta_correcciones = rutas.resolver(
        configuracion.obtener("informe.correcciones",
                              "config/informe_correcciones.yaml"), base)
    correcciones = leer_correcciones(ruta_correcciones)
    sustituciones = leer_sustituciones(ruta_correcciones)
    decimales = int(configuracion.obtener("informe.decimales", 2))

    resultado = ResultadoM15(plantilla=str(origen))

    registro.registrar_cabecera(
        logger, MODULO, DESCRIPCION, config=configuracion,
        insumos={"plantilla": str(origen),
                 "figuras": rutas.relativa(graficos, base),
                 "tablas declaradas": str(len(declaracion))},
        parametros={"informe.archivo": configuracion.obtener("informe.archivo")},
    )

    try:
        documento = plantilla_docx.abrir(origen)
    except (plantilla_docx.ErrorPlantilla, ErrorRutas) as error:
        resultado.hallazgos.append(Hallazgo(
            BLOQUEANTE, "informe.plantilla", str(error)))
        return _cerrar(logger, resultado, base, ruta_json, inicio,
                       SALIDA_BLOQUEANTE)

    ancho = _ancho_util(documento)
    cuerpo = list(documento.element.body)
    parrafos = {p._element: p for p in documento.paragraphs}
    tablas = {t._tbl: t for t in documento.tables}
    plan, ambiguas, sin_uso = planear_correcciones(documento, correcciones)
    resultado.correcciones_ambiguas = ambiguas
    resultado.correcciones_sin_uso = sin_uso

    with registro.bloque(logger, "Instrucciones de la plantilla"):
        for posicion, elemento in enumerate(cuerpo):
            parrafo = parrafos.get(elemento)
            if parrafo is None:
                continue
            tipo, argumento = clasificar(parrafo.text)
            if tipo == "analisis":
                resultado.analisis_pendientes += 1
            elif tipo == "figura":
                _resolver_figura(parrafo, argumento, [graficos, individuales],
                                 ancho, resultado, base, plan)
            elif tipo == "tabla":
                _resolver_tabla(cuerpo, posicion, tablas, parrafos,
                                argumento, declaracion, delimitador, base,
                                resultado, decimales, sustituciones)

        # LAS INSTRUCCIONES TAMBIEN VIVEN DENTRO DE LAS TABLAS. La plantilla
        # compone algunas figuras en celdas, para ponerlas de a dos, y esos
        # parrafos no estan en documento.paragraphs: son 16 de las 92, y sin
        # recorrerlas quedaban mudas sin que nada lo dijera.
        for tabla_contenedora in documento.tables:
            for fila in tabla_contenedora.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        tipo, argumento = clasificar(parrafo.text)
                        if tipo == "analisis":
                            resultado.analisis_pendientes += 1
                        elif tipo == "figura":
                            _resolver_figura(
                                parrafo, argumento, [graficos, individuales],
                                _ancho_de_celda(celda, ancho), resultado, base,
                                plan)

        logger.info("%d figura(s) puestas, %d tabla(s) llenadas, %d analisis "
                    "sin tocar", len(resultado.figuras_puestas),
                    len(resultado.tablas_llenadas),
                    resultado.analisis_pendientes)

    marcados = marcar_campos_para_actualizar(documento)
    logger.info("%d campo(s) marcados para que Word los actualice al abrir",
                marcados)

    salida = (Path(ruta_salida) if ruta_salida is not None
              else rutas.directorio("resultados", base, crear=True) / str(
                  configuracion.obtener("informe.archivo")))
    salida.parent.mkdir(parents=True, exist_ok=True)
    try:
        documento.save(str(salida))
    except PermissionError:
        # WORD BLOQUEA EL ARCHIVO MIENTRAS LO TIENE ABIERTO. Sin este aviso, el
        # modulo muere con una traza de la libreria que no dice que hay que
        # cerrar el documento.
        resultado.hallazgos.append(Hallazgo(
            BLOQUEANTE, "informe.archivo_bloqueado",
            f"no se pudo escribir {salida.name}: el sistema lo tiene abierto, "
            "casi siempre porque el documento esta abierto en Word. Cerrarlo y "
            "repetir, o escribir en otro nombre con --salida.",
        ))
        return _cerrar(logger, resultado, base, ruta_json, inicio,
                       SALIDA_BLOQUEANTE)
    resultado.documento = rutas.relativa(salida, base)
    resultado.productos.append(resultado.documento)

    resultado.hallazgos.extend(_resumir(resultado))
    codigo = (SALIDA_BLOQUEANTE
              if any(h.severidad == BLOQUEANTE for h in resultado.hallazgos)
              else SALIDA_CORRECTA)
    return _cerrar(logger, resultado, base, ruta_json, inicio, codigo)


def _resolver_figura(parrafo, nombre, raices, ancho, resultado, base,
                     plan=None) -> None:
    """Pone la figura si existe; si no, deja la instrucción y lo reporta."""
    entrada = (plan or {}).get(parrafo._element)
    if entrada is not None and str(entrada.get("archivo", "")).strip():
        # LA CORRECCION SE REGISTRA Y NO SE APLICA EN SILENCIO. Cambiar la
        # figura que la plantilla pide es una decision con margen, y un estudio
        # que no puede explicar sus cambios no es defendible.
        resultado.figuras_corregidas.append({
            "leyenda": str(entrada.get("leyenda", "")), "pedia": nombre,
            "archivo": str(entrada["archivo"]),
            "motivo": str(entrada.get("motivo", "")).strip()})
        nombre = str(entrada["archivo"])
    ruta = buscar_figura(nombre, raices)
    if ruta is None:
        # LA INSTRUCCION SE DEJA INTACTA. Borrarla dejaria un hueco mudo en el
        # informe y nadie sabria que falta una figura.
        resultado.figuras_ausentes.append({"archivo": nombre})
        return
    poner_figura(parrafo, ruta, ancho)
    resultado.figuras_puestas.append(rutas.relativa(ruta, base))


def _resolver_tabla(cuerpo, posicion, tablas, parrafos, numero, declaracion,
                    delimitador, base, resultado, decimales=2,
                    sustituciones=()) -> None:
    """
    Llena la primera tabla que sigue a la instrucción, buscada por su leyenda.

    ENTRE LA INSTRUCCION Y LA TABLA VA LA LEYENDA, de modo que no basta con
    mirar el elemento siguiente. Se avanza hasta encontrar la tabla, y de paso
    se recoge esa leyenda, que es la clave con la que se busca la fuente: el
    número de la instrucción no identifica a la tabla (ver
    'leer_declaracion_tablas').
    """
    tabla, leyenda = None, ""
    for elemento in cuerpo[posicion + 1:posicion + 8]:
        if elemento in tablas:
            tabla = tablas[elemento]
            break
        parrafo = parrafos.get(elemento)
        if parrafo is not None and not leyenda and parrafo.text.strip():
            leyenda = parrafo.text.strip()
    if tabla is None:
        resultado.tablas_sin_fuente.append(f"{numero} (sin tabla detras)")
        return

    entrada = declaracion.get(_normalizar_leyenda(leyenda))
    if entrada is None:
        resultado.tablas_sin_fuente.append(leyenda or numero)
        return

    # EL ANCHO DECLARADO TIENE QUE SER EL DE LA TABLA. Si se declaran menos
    # columnas de las que la tabla tiene, las sobrantes se quedan con los
    # numeros del estudio del que se copio la plantilla, y la tabla sale
    # mezclando dos estudios sin que nada lo advierta.
    matriz = entrada.get("matriz")
    if matriz:
        columnas = [str(matriz.get("fila", ""))] + [
            str(v) for v in matriz.get("orden") or []]
    else:
        columnas = [str(c) for c in entrada.get("columnas") or []]
    crece = bool((matriz or entrada).get("crecer_columnas"))
    if len(columnas) != len(tabla.columns) and not crece:
        resultado.tablas_sin_fuente.append(
            f"{leyenda or numero}: se declararon {len(columnas)} columna(s) y "
            f"la tabla tiene {len(tabla.columns)}")
        return
    if crece and len(columnas) < len(tabla.columns):
        # Crecer no es encoger. Si la tabla tiene MAS columnas de las
        # declaradas, las sobrantes se quedarian con lo que la plantilla traia.
        resultado.tablas_sin_fuente.append(
            f"{leyenda or numero}: se declararon {len(columnas)} columna(s) y "
            f"la tabla ya tiene {len(tabla.columns)}; 'crecer_columnas' "
            "ensancha, no estrecha")
        return

    try:
        filas = leer_tabla_csv(rutas.resolver(str(entrada["fuente"]), base),
                               delimitador)
        filtro = entrada.get("filtro") or {}
        if filtro:
            # UNA FUENTE PUEDE TRAER MAS DE LO QUE LA TABLA PIDE. La de
            # precipitacion por subcuenca viaja etiquetada por hipotesis y
            # escenario, y sin acotarla la matriz mezclaria escenarios en la
            # misma celda quedandose con el ultimo.
            filas = [f for f in filas
                     if all(str(f.get(k, "")) == str(v)
                            for k, v in filtro.items())]
            if not filas:
                raise ErrorFormato(
                    f"el filtro {filtro} no dejo ninguna fila de la fuente.")
        if matriz:
            detalle = llenar_matriz(tabla, filas, matriz,
                                    int(entrada.get("encabezados", 1)),
                                    decimales, sustituciones)
        else:
            detalle = llenar_tabla(tabla, filas, columnas,
                                   int(entrada.get("encabezados", 1)),
                                   decimales, sustituciones,
                                   entrada.get("titulos") or ())
    except (ErrorRutas, ErrorFormato, KeyError, ValueError) as error:
        resultado.tablas_sin_fuente.append(f"{leyenda or numero} ({error})")
        return
    # LA INSTRUCCION SE BORRA DEL INFORME, NUNCA DE LA PLANTILLA. Va marcada en
    # verde y el consultor la usa para saber que falta; una vez la tabla lleva
    # los datos del estudio, dejarla es dejar en el entregable una nota que
    # dice que la tabla esta sin llenar. Con las figuras ya ocurria, porque la
    # imagen sustituye al texto; las 27 de tabla sobrevivian.
    instruccion = parrafos.get(cuerpo[posicion])
    if instruccion is not None:
        instruccion._element.getparent().remove(instruccion._element)

    detalle["numero"] = numero
    detalle["leyenda"] = leyenda
    detalle["fuente"] = str(entrada["fuente"])
    resultado.tablas_llenadas.append(detalle)


def _resumir(resultado: ResultadoM15) -> list[Hallazgo]:
    """Lo que el módulo resolvió y lo que queda por hacer a mano."""
    hallazgos: list[Hallazgo] = []

    hallazgos.append(Hallazgo(
        INFORMATIVO, "informe.resuelto",
        f"{len(resultado.figuras_puestas)} figura(s) insertadas y "
        f"{len(resultado.tablas_llenadas)} tabla(s) llenadas con los datos de "
        "este estudio. Se editó el documento, no se reconstruyó: las "
        "referencias cruzadas, las leyendas numeradas por campo, los "
        "hipervínculos y los marcadores quedan intactos.",
    ))

    if resultado.analisis_pendientes:
        hallazgos.append(Hallazgo(
            INFORMATIVO, "informe.cambio_de_modelo",
            "AQUI TERMINA LO MECANICO Y EMPIEZA EL JUICIO. Lo hecho hasta este "
            "punto (figuras, tablas, mapas, anexos) es determinista: se ejecuta "
            "y se comprueba. Lo que queda exige mirar el resultado y decir que "
            "significa, defenderlo ante una interventoria y reconocer cuando un "
            "numero no cuadra. Si se trabaja con asistencia de un modelo, es el "
            "momento de pasar al de mayor capacidad de razonamiento: en el resto "
            "de la cadena la diferencia no se nota, y aqui es toda la "
            "diferencia."))
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.analisis_pendiente",
            f"{resultado.analisis_pendientes} instruccion(es) de análisis "
            "siguen en verde y sin resolver, que es lo previsto: exigen mirar "
            "el resultado y decir qué significa, y eso no se programa. El "
            "informe NO está terminado hasta que se redacten y se borren.",
        ))

    if resultado.figuras_corregidas:
        detalle = "; ".join(
            f"bajo '{c['leyenda']}' se puso {c['archivo']} y no {c['pedia']}"
            for c in resultado.figuras_corregidas)
        hallazgos.append(Hallazgo(
            INFORMATIVO, "informe.figuras_corregidas",
            f"{len(resultado.figuras_corregidas)} instruccion(es) de la "
            f"plantilla pedían la figura equivocada y se corrigieron según "
            f"informe.correcciones: {detalle}. La plantilla original conserva "
            "el error; el motivo de cada cambio está declarado ahí.",
        ))

    if resultado.correcciones_ambiguas:
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.correcciones_ambiguas",
            f"{len(resultado.correcciones_ambiguas)} corrección(es) declaradas "
            f"emparejan con más de una instruccion y NO se aplicaron a ninguna: "
            f"{resultado.correcciones_ambiguas}. La plantilla repite esa "
            "leyenda en figuras distintas. Añada 'pedia' con el archivo que la "
            "instruccion nombra hoy para desempatar.",
        ))

    if resultado.correcciones_sin_uso:
        # UNA CORRECCION QUE NO EMPAREJA NO HACE NADA Y NADIE LO NOTA. O la
        # leyenda esta mal escrita, o el consultor ya corrigio su plantilla y
        # la entrada sobra.
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.correcciones_sin_uso",
            f"{len(resultado.correcciones_sin_uso)} corrección(es) declaradas "
            f"no emparejaron con ninguna leyenda de la plantilla: "
            f"{resultado.correcciones_sin_uso}. O la leyenda está mal escrita "
            "y la corrección no se aplicó, o la plantilla ya se corrigió y la "
            "entrada sobra.",
        ))

    if resultado.figuras_ausentes:
        nombres = sorted({f["archivo"] for f in resultado.figuras_ausentes})
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.figuras_ausentes",
            f"{len(nombres)} figura(s) que la plantilla pide no están en el "
            f"estudio: {nombres}. Su instrucción se dejó en el documento en "
            "lugar de borrarla: un hueco mudo no se ve al revisar.",
        ))

    if resultado.tablas_sin_fuente:
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.tablas_sin_fuente",
            f"{len(resultado.tablas_sin_fuente)} tabla(s) quedaron con los "
            f"datos de la plantilla: {sorted(resultado.tablas_sin_fuente)}. "
            "Falta declarar su fuente en informe.tablas, o la fuente declarada "
            "no se pudo leer.",
        ))

    sin_columna = sorted({c for t in resultado.tablas_llenadas
                          for c in t.get("sin_columna", [])})
    if sin_columna:
        hallazgos.append(Hallazgo(
            ADVERTENCIA, "informe.columnas_ausentes",
            f"se declararon columnas que su CSV no tiene: {sin_columna}. Esas "
            "celdas quedaron VACIAS, no con el valor anterior: un dato del "
            "estudio pasado en una tabla de este es peor que un hueco.",
        ))

    hallazgos.append(Hallazgo(
        INFORMATIVO, "informe.actualizar_campos",
        "los campos del documento (las leyendas numeradas y los índices de "
        "contenido, de ilustraciones y de tablas) quedan MARCADOS para que "
        "Word los recalcule al abrirlo, de modo que las leyendas nuevas entran "
        "en la numeración de su capítulo sin hacer nada. Si el visor no lo "
        "honra, y LibreOffice y las vistas previas web no lo hacen, se ven los "
        "números en caché: abrirlo en Word, o pulsar Ctrl+E y F9.",
    ))
    return hallazgos


def _cerrar(logger, resultado, base, ruta_json, inicio, codigo):
    """Emite el reporte, escribe el JSON y cierra el log."""
    orden = {BLOQUEANTE: 0, ADVERTENCIA: 1, INFORMATIVO: 2}
    hallazgos = sorted(resultado.hallazgos,
                       key=lambda h: (orden.get(h.severidad, 9), h.clave))

    logger.info(registro.SEPARADOR)
    for severidad, emitir in ((BLOQUEANTE, logger.error),
                              (ADVERTENCIA, logger.warning),
                              (INFORMATIVO, logger.info)):
        grupo = [h for h in hallazgos if h.severidad == severidad]
        if not grupo:
            continue
        emitir("%s (%d)", severidad, len(grupo))
        for hallazgo in grupo:
            emitir("  %-40s %s", hallazgo.clave, hallazgo.mensaje)

    conteo = esquema.resumen_por_severidad(hallazgos)
    logger.info(registro.SEPARADOR)
    logger.info("RESUMEN: %d bloqueante(s), %d advertencia(s), %d informativo(s)",
                conteo[BLOQUEANTE], conteo[ADVERTENCIA], conteo[INFORMATIVO])

    if ruta_json is None:
        ruta_json = (rutas.directorio("procesado", base, crear=True)
                     / "M15_informe.json")
    reporte = {
        "modulo": MODULO,
        "plantilla": resultado.plantilla,
        "documento": resultado.documento,
        "figuras_puestas": resultado.figuras_puestas,
        "figuras_ausentes": resultado.figuras_ausentes,
        "tablas_llenadas": resultado.tablas_llenadas,
        "tablas_sin_fuente": resultado.tablas_sin_fuente,
        "analisis_pendientes": resultado.analisis_pendientes,
        "productos": resultado.productos,
        "resumen": conteo,
        "codigo_salida": codigo,
        "conforme": codigo == SALIDA_CORRECTA,
        "hallazgos": [h.como_dict() for h in hallazgos],
    }
    ruta_json.parent.mkdir(parents=True, exist_ok=True)
    ruta_json.write_text(json.dumps(reporte, ensure_ascii=False, indent=2),
                         encoding="utf-8")

    productos = {f"producto {i}": p
                 for i, p in enumerate(resultado.productos, start=1)}
    productos["reporte JSON"] = rutas.relativa(ruta_json, base)
    archivo_log = registro.ruta_log(logger)
    if archivo_log is not None:
        productos["log de ejecucion"] = rutas.relativa(archivo_log, base)

    registro.registrar_cierre(
        logger, MODULO, "CORRECTO" if codigo == SALIDA_CORRECTA else "DETENIDO",
        segundos=time.perf_counter() - inicio, productos=productos)
    return codigo, hallazgos


# =============================================================================
# Interfaz de linea de comandos
# =============================================================================
def main(argv=None) -> int:
    """Punto de entrada. Devuelve el codigo de salida del proceso."""
    analizador = argparse.ArgumentParser(
        prog="M15_informe.py",
        description="Resuelve figuras y tablas sobre la plantilla del informe.")
    analizador.add_argument("--raiz", type=Path, default=None)
    analizador.add_argument("--config", type=Path, default=None)
    analizador.add_argument("--plantilla", type=Path, default=None)
    analizador.add_argument("--salida", type=Path, default=None,
                            help="escribe en otro archivo, util si Word tiene "
                                 "abierto el habitual")
    analizador.add_argument("--json", type=Path, default=None,
                            dest="json_salida")
    analizador.add_argument("--silencioso", action="store_true")
    argumentos = analizador.parse_args(argv)
    try:
        codigo, _ = ejecutar(
            raiz=argumentos.raiz, ruta_config=argumentos.config,
            ruta_json=argumentos.json_salida,
            ruta_plantilla=argumentos.plantilla,
            ruta_salida=argumentos.salida,
            consola=not argumentos.silencioso,
        )
        return codigo
    except (ErrorRutas, ErrorConfiguracion, ErrorFormato) as exc:
        print(f"ERROR    | {exc}", file=sys.stderr)
        return SALIDA_ERROR
    except ErrorHidrologia as exc:
        print(f"ERROR    | {exc}", file=sys.stderr)
        return SALIDA_ERROR


if __name__ == "__main__":
    sys.exit(main())
