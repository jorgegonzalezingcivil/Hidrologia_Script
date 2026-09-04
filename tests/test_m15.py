# -*- coding: utf-8 -*-
"""
Pruebas del M15: resolución de las instrucciones de la plantilla.

Se prueban las funciones puras y la declaración real del repositorio. La
edición del documento exige python-docx y una plantilla de 3,7 MB, de modo que
se verifica sobre la del estudio en las pruebas de integración y aquí solo lo
que decide qué se hace con cada instrucción.

    python tests/test_m15.py
"""

from __future__ import annotations

import sys
import tempfile
import unittest
from pathlib import Path

_RAIZ_REPO = Path(__file__).resolve().parents[1]
_DIRECTORIO_SRC = _RAIZ_REPO / "src"
if str(_DIRECTORIO_SRC) not in sys.path:
    sys.path.insert(0, str(_DIRECTORIO_SRC))

import M15_informe as m15  # noqa: E402
from comun.errores import ErrorRutas  # noqa: E402


class PruebaClasificacion(unittest.TestCase):
    """
    SE DECIDE POR EL TEXTO Y NO POR EL COLOR. El sombreado se pierde al copiar y
    pegar un párrafo, y entonces la instrucción quedaría muda sin que nada lo
    señalara.
    """

    def test_reconoce_una_figura_y_extrae_su_archivo(self) -> None:
        tipo, argumento = m15.clasificar("Colocar Figura: M10_mapa_cn.png")
        self.assertEqual(tipo, "figura")
        self.assertEqual(argumento, "M10_mapa_cn.png")

    def test_la_carpeta_que_la_instruccion_declara_viaja_con_el_archivo(self) -> None:
        """
        Antes se descartaba, y eso puso tres figuras equivocadas en el informe.

        'compuesto.png', 'nina.png' y 'nino.png' existen a la vez en
        'individuales/isoyetas_fase', que son los campos de precipitación, y en
        'individuales/contraste_fases', que son los mapas de anomalía. La
        búsqueda por nombre devolvía la primera que encontraba y el informe
        mostraba un mapa de cambio porcentual bajo una leyenda que anuncia
        precipitación total. La cuarta, 'neutral.png', salía bien por
        casualidad: solo existe en una de las dos carpetas.

        La instrucción SÍ dice de qué carpeta es. Lo que faltaba era leerlo.
        """
        tipo, argumento = m15.clasificar(
            "Colocar Figura: compuesto.png de la carpeta individuales/isoyetas_fase")
        self.assertEqual(tipo, "figura")
        self.assertEqual(argumento,
                         "compuesto.png|individuales/isoyetas_fase")

    def test_sin_carpeta_declarada_el_argumento_es_solo_el_archivo(self) -> None:
        self.assertEqual(m15.clasificar("Colocar Figura: M10_mapa_cn.png")[1],
                         "M10_mapa_cn.png")

    def test_reconoce_una_tabla_y_su_numero(self) -> None:
        tipo, argumento = m15.clasificar(
            "Completar la Tabla 2-1 con datos según estaciones identificadas")
        self.assertEqual(tipo, "tabla")
        self.assertEqual(m15._normalizar_numero(argumento), "21")

    def test_reconoce_el_analisis_y_no_lo_resuelve(self) -> None:
        tipo, _argumento = m15.clasificar("Analizar Ilustración 3-5 y 3-6.")
        self.assertEqual(tipo, "analisis")

    def test_un_parrafo_corriente_no_es_instruccion(self) -> None:
        for texto in ("En la Tabla 2-1 se presentan las características",
                      "Fuente: IDEAM, 2024.", "", "   ",
                      "Ilustración 22. Localización Estaciones"):
            self.assertEqual(m15.clasificar(texto)[0], "")

    def test_no_distingue_mayusculas(self) -> None:
        self.assertEqual(m15.clasificar("colocar figura: x.png")[0], "figura")
        self.assertEqual(m15.clasificar("COMPLETAR LA TABLA 3-1")[0], "tabla")


class PruebaNumeroDeTabla(unittest.TestCase):
    """
    Word compone la leyenda con campos y el guion se pierde al extraer su texto,
    o aparece como guion corto, largo o de no separación según cómo se escribió.
    """

    def test_los_tres_guiones_dan_el_mismo_numero(self) -> None:
        self.assertEqual(m15._normalizar_numero("2-1"), "21")
        self.assertEqual(m15._normalizar_numero("2‑1"), "21")
        self.assertEqual(m15._normalizar_numero("2—1"), "21")
        self.assertEqual(m15._normalizar_numero("21"), "21")

    def test_distingue_numeros_distintos(self) -> None:
        # 3-1 y 3-10 no pueden colapsar en el mismo, o una tabla se llenaria
        # con los datos de la otra.
        self.assertNotEqual(m15._normalizar_numero("3-1"),
                            m15._normalizar_numero("3-10"))


class PruebaBusquedaDeFiguras(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = Path(tempfile.mkdtemp())
        (self.tmp / "individuales" / "enso").mkdir(parents=True)
        (self.tmp / "M10_mapa_cn.png").write_bytes(b"x")
        (self.tmp / "individuales" / "enso" / "guasca.png").write_bytes(b"x")

    def test_encuentra_en_la_raiz(self) -> None:
        self.assertIsNotNone(m15.buscar_figura("M10_mapa_cn.png", [self.tmp]))

    def test_encuentra_en_una_subcarpeta(self) -> None:
        # El consultor no tiene por que saber en que tema la dejo cada modulo.
        encontrada = m15.buscar_figura("guasca.png", [self.tmp])
        self.assertIsNotNone(encontrada)
        self.assertEqual(encontrada.name, "guasca.png")

    def test_lo_que_no_existe_devuelve_nada(self) -> None:
        self.assertIsNone(m15.buscar_figura("no_existe.png", [self.tmp]))

    def test_una_raiz_ausente_no_es_error(self) -> None:
        self.assertIsNone(
            m15.buscar_figura("x.png", [self.tmp / "no_esta"]))

    def test_la_carpeta_declarada_decide_entre_dos_del_mismo_nombre(self) -> None:
        """
        El nombre NO es unico, y darlo por unico costo tres figuras
        equivocadas en el entregable.
        """
        (self.tmp / "individuales" / "otra").mkdir(parents=True)
        (self.tmp / "individuales" / "enso" / "repetida.png").write_bytes(b"a")
        (self.tmp / "individuales" / "otra" / "repetida.png").write_bytes(b"b")

        encontrada = m15.buscar_figura("repetida.png", [self.tmp],
                                       "individuales/otra")
        self.assertEqual(encontrada.read_bytes(), b"b")
        self.assertEqual(encontrada.parent.name, "otra")

    def test_la_carpeta_vale_con_el_prefijo_y_sin_el(self) -> None:
        # La instruccion escribe 'individuales/enso' y la raiz de busqueda
        # puede ser ya el propio 'individuales'.
        raiz = self.tmp / "individuales"
        self.assertIsNotNone(
            m15.buscar_figura("guasca.png", [raiz], "individuales/enso"))
        self.assertIsNotNone(
            m15.buscar_figura("guasca.png", [raiz], "enso"))

    def test_las_candidatas_no_se_cuentan_dos_veces(self) -> None:
        # Las raices estan anidadas: 'individuales' cuelga del directorio de
        # graficos. Sin deduplicar, ocho hietogramas se reportaban como
        # ambiguos sin serlo, y un aviso que casi siempre sobra se ignora.
        candidatas = m15.figuras_ambiguas(
            "guasca.png", [self.tmp, self.tmp / "individuales"])
        self.assertEqual(len(candidatas), 1)


class PruebaDeclaracionDeTablas(unittest.TestCase):
    """La declaración real del repositorio, no una inventada para la prueba."""

    def setUp(self) -> None:
        self.declaracion = m15.leer_declaracion_tablas(
            _RAIZ_REPO / "config" / "informe_tablas.yaml")

    def test_la_declaracion_del_repositorio_se_lee(self) -> None:
        self.assertTrue(self.declaracion)

    def test_toda_tabla_declara_fuente_y_como_se_llena(self) -> None:
        # 'columnas' para una lista, 'matriz' para una matriz, y una de las dos
        # siempre: sin ninguna, la tabla se quedaria con lo que la plantilla
        # traia y el modulo diria que la lleno.
        for clave, entrada in self.declaracion.items():
            self.assertTrue(entrada.get("fuente"), clave)
            self.assertTrue(entrada.get("columnas") or entrada.get("matriz"),
                            clave)

    def test_se_indexa_por_leyenda_y_no_por_numero(self) -> None:
        # El numero de la instruccion es el del informe del que se copio el
        # apartado: en esta plantilla hay dos que dicen 5-1 delante de tablas
        # distintas. La clave tiene que ser la leyenda.
        for clave, entrada in self.declaracion.items():
            self.assertEqual(clave,
                             m15._normalizar_leyenda(str(entrada["leyenda"])))

    def test_las_leyendas_no_colisionan(self) -> None:
        # Dos tablas con la misma leyenda normalizada harian que una se llenara
        # con los datos de la otra sin ninguna senal.
        leyendas = [str(e.get("leyenda")) for e in self.declaracion.values()]
        self.assertEqual(len(leyendas),
                         len({m15._normalizar_leyenda(l) for l in leyendas}))

    def test_toda_leyenda_declarada_existe_en_la_plantilla(self) -> None:
        # Una leyenda mal escrita no emparejaria con nada y la tabla quedaria
        # sin llenar, que es justo lo que no se nota al revisar el informe.
        plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not plantilla.is_file():
            self.skipTest("la plantilla saneada no esta en el repositorio")
        import docx_plantilla as dp
        documento = dp.abrir(plantilla)
        presentes = {m15._normalizar_leyenda(p.text)
                     for p in documento.paragraphs if p.text.strip()}
        for clave in self.declaracion:
            self.assertIn(clave, presentes)

    def test_los_encabezados_son_al_menos_uno(self) -> None:
        # Con cero, la primera fila de titulos se sobreescribiria con datos.
        for clave, entrada in self.declaracion.items():
            self.assertGreaterEqual(int(entrada.get("encabezados", 1)), 1,
                                    clave)

    def test_una_declaracion_ausente_no_es_error(self) -> None:
        # Significa que aun no se declaro ninguna tabla, no que falte algo.
        self.assertEqual(
            m15.leer_declaracion_tablas(Path("no_existe.yaml")), {})


class PruebaNormalizacionDeLeyenda(unittest.TestCase):
    """
    Las leyendas de la plantilla dicen 'Tabla -.' porque sus campos SEQ no
    tienen resultado en cache. El prefijo es justamente lo que no se compara.
    """

    def test_el_prefijo_numerado_no_cuenta(self) -> None:
        esperado = m15._normalizar_leyenda("Perímetro microcuencas")
        for texto in ("Tabla -. Perímetro microcuencas",
                      "Tabla 3-2. Perímetro microcuencas",
                      "Tabla 5-10. Perímetro microcuencas"):
            self.assertEqual(m15._normalizar_leyenda(texto), esperado)

    def test_las_tildes_y_las_mayusculas_no_cuentan(self) -> None:
        # La declaracion la escribe el consultor a mano.
        self.assertEqual(m15._normalizar_leyenda("Área  microcuencas"),
                         m15._normalizar_leyenda("area microcuencas"))

    def test_distingue_tablas_distintas(self) -> None:
        self.assertNotEqual(
            m15._normalizar_leyenda("Tabla -. Coeficiente de forma microcuencas"),
            m15._normalizar_leyenda("Tabla -. Coeficiente de compacidad microcuencas"))

class PruebaCorrecciones(unittest.TestCase):
    """
    La plantilla pide figuras equivocadas y el original no se toca: la
    correccion se declara, y tiene que emparejar con UNA instruccion.
    """

    def setUp(self) -> None:
        self.plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not self.plantilla.is_file():
            self.skipTest("la plantilla saneada no esta en el repositorio")
        import docx_plantilla as dp
        self.documento = dp.abrir(self.plantilla)
        self.correcciones = m15.leer_correcciones(
            _RAIZ_REPO / "config" / "informe_correcciones.yaml")

    def test_estar_vacia_es_lo_esperado(self) -> None:
        # Las dos correcciones que hubo estan CONSOLIDADAS EN LA PLANTILLA, por
        # tools/consolidar_plantilla.py: el repositorio es la fuente de verdad y
        # el defecto se arreglo en su sitio. El mecanismo se conserva para la
        # proxima, y lo que se comprueba abajo es que si hay alguna declarada
        # este completa y empareje con una sola instruccion.
        self.assertIsInstance(self.correcciones, dict)

    def test_toda_correccion_declara_archivo_y_motivo(self) -> None:
        for clave, entrada in self.correcciones.items():
            self.assertTrue(entrada.get("archivo"), clave)
            self.assertTrue(entrada.get("motivo"), clave)

    def test_ninguna_correccion_queda_ambigua_ni_sin_uso(self) -> None:
        # Una que empareja con varias sustituiria figuras que estaban bien; una
        # que no empareja con ninguna no hace nada y nadie lo nota.
        _plan, ambiguas, sin_uso = m15.planear_correcciones(
            self.documento, self.correcciones)
        self.assertEqual(ambiguas, [])
        self.assertEqual(sin_uso, [])

    def test_cada_correccion_apunta_a_una_sola_instruccion(self) -> None:
        plan, _a, _s = m15.planear_correcciones(
            self.documento, self.correcciones)
        self.assertEqual(len(plan), len(self.correcciones))

    def test_sin_desempate_la_leyenda_repetida_es_ambigua(self) -> None:
        """
        Una leyenda que encabeza dos instrucciones no identifica ninguna.

        SE COMPRUEBA SOBRE UN DOCUMENTO PROPIO Y NO SOBRE LA PLANTILLA. Antes
        se apoyaba en que 'Areas microcuencas' encabezaba tres instrucciones de
        la plantilla, y al corregir dos de esas tres parejas mal emparejadas la
        prueba dejo de comprobar nada: su premisa era un defecto del documento,
        no una propiedad del mecanismo.
        """
        import docx

        documento = docx.Document()
        for _ in range(2):
            documento.add_paragraph("Ilustración 1-1. Leyenda repetida",
                                    style="Caption")
            documento.add_paragraph("Colocar Figura: vieja.png")

        suelta = {m15._normalizar_leyenda("Leyenda repetida"): {
            "leyenda": "Leyenda repetida", "archivo": "x.png"}}
        plan, ambiguas, _s = m15.planear_correcciones(documento, suelta)
        self.assertEqual(plan, {})
        self.assertEqual(len(ambiguas), 1)

    def test_una_declaracion_ausente_no_es_error(self) -> None:
        self.assertEqual(m15.leer_correcciones(Path("no_existe.yaml")), {})


class PruebaFormatoDeCifras(unittest.TestCase):
    """
    El consultor pidio dos decimales en toda cifra del informe. La regla no
    puede alcanzar a lo que no es una medida.
    """

    def test_una_medida_queda_con_dos_decimales(self) -> None:
        self.assertEqual(m15.formatear_numero("45.2", 2), "45.20")
        self.assertEqual(m15.formatear_numero("117.2345", 2), "117.23")

    def test_un_anio_o_un_codigo_no_se_tocan(self) -> None:
        # Redondear un entero lo convertiria en '1983.00' y en
        # '21201230.00', que no son ni un anio ni un codigo de estacion.
        for texto in ("1983", "21201230", "20011115", "12"):
            self.assertEqual(m15.formatear_numero(texto, 2), texto)

    def test_lo_que_no_es_numero_pasa_intacto(self) -> None:
        for texto in ("lognormal2", "CO", "", "Q. NN", "2.33 anios"):
            self.assertEqual(m15.formatear_numero(texto, 2), texto)

    def test_una_cifra_pequena_no_se_convierte_en_cero(self) -> None:
        # Un caudal de 0.0004 m3/s es un dato; '0.00' es una perdida muda.
        self.assertEqual(m15.formatear_numero("0.0004", 2), "0.0004")
        self.assertEqual(m15.formatear_numero("0.0", 2), "0.00")


class PruebaSustitucionDeNombres(unittest.TestCase):
    """El catalogo entrega la razon social; el informe usa la sigla."""

    def setUp(self) -> None:
        self.reglas = m15.leer_sustituciones(
            _RAIZ_REPO / "config" / "informe_correcciones.yaml")

    def test_la_declaracion_del_repositorio_se_lee(self) -> None:
        self.assertTrue(self.reglas)

    def test_toda_regla_declara_que_busca_y_por_que(self) -> None:
        for regla in self.reglas:
            self.assertTrue(regla.get("busca"))
            self.assertTrue(regla.get("pone"))
            self.assertTrue(regla.get("motivo"))

    def test_el_nombre_largo_del_ideam_queda_en_sigla(self) -> None:
        largo = "INSTITUTO DE HIDROLOGÍA METEOROLOGÍA Y ESTUDIOS AMBIENTALES"
        self.assertEqual(m15.aplicar_sustituciones(largo, self.reglas), "IDEAM")

    def test_lo_que_no_coincide_pasa_intacto(self) -> None:
        self.assertEqual(
            m15.aplicar_sustituciones("EAAB", self.reglas), "EAAB")

    def test_sin_reglas_no_cambia_nada(self) -> None:
        self.assertEqual(m15.aplicar_sustituciones("IDEAM", []), "IDEAM")


class PruebaAnchoDeclarado(unittest.TestCase):
    """
    EL ANCHO DECLARADO TIENE QUE SER EL DE LA TABLA. Si sobran columnas en la
    tabla, se quedan con lo que la plantilla traia y el informe mezcla dos
    estudios sin que nada lo advierta.
    """

    def setUp(self) -> None:
        plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not plantilla.is_file():
            self.skipTest("la plantilla saneada no esta en el repositorio")
        import docx_plantilla as dp
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        documento = dp.abrir(plantilla)
        elementos = []
        for hijo in documento.element.body.iterchildren():
            if hijo.tag.endswith("}p"):
                elementos.append(("p", Paragraph(hijo, documento)))
            elif hijo.tag.endswith("}tbl"):
                elementos.append(("t", Table(hijo, documento)))
        self.tablas = {}
        for i, (clase, obj) in enumerate(elementos):
            if clase != "p" or m15.clasificar(obj.text)[0] != "tabla":
                continue
            leyenda, tabla = "", None
            for j in range(i + 1, min(len(elementos), i + 8)):
                if elementos[j][0] == "t":
                    tabla = elementos[j][1]
                    break
                texto = elementos[j][1].text.strip()
                if texto and not leyenda:
                    leyenda = texto
            if tabla is not None:
                self.tablas[m15._normalizar_leyenda(leyenda)] = tabla
        self.declaracion = m15.leer_declaracion_tablas(
            _RAIZ_REPO / "config" / "informe_tablas.yaml")

    def test_cada_declaracion_cuadra_con_su_tabla(self) -> None:
        for clave, entrada in self.declaracion.items():
            tabla = self.tablas.get(clave)
            self.assertIsNotNone(tabla, clave)
            matriz = entrada.get("matriz")
            if matriz:
                ancho = 1 + len(matriz.get("orden") or [])
            else:
                ancho = len(entrada.get("columnas") or [])
            if (matriz or entrada).get("crecer_columnas"):
                # La tabla se ensancha al llenarla: la plantilla viene
                # dimensionada para cuatro microcuencas. Crecer no es encoger.
                self.assertGreaterEqual(ancho, len(tabla.columns), clave)
            else:
                self.assertEqual(ancho, len(tabla.columns), clave)

    def test_la_que_se_ensancha_declara_sus_titulos(self) -> None:
        # Una columna anadida sale sin titulo, y una tabla con titulos en
        # blanco no se entiende. En modo matriz los titulos son el 'orden'.
        for clave, entrada in self.declaracion.items():
            if not entrada.get("crecer_columnas"):
                continue
            titulos = entrada.get("titulos") or []
            self.assertEqual(len(titulos), len(entrada.get("columnas") or []),
                             clave)

    def test_los_encabezados_dejan_al_menos_una_fila_de_datos(self) -> None:
        for clave, entrada in self.declaracion.items():
            tabla = self.tablas[clave]
            self.assertGreater(len(tabla.rows),
                               int(entrada.get("encabezados", 1)), clave)


class PruebaDeclaracionDeMatrices(unittest.TestCase):
    """Doce tablas del informe son matrices y no listas."""

    def setUp(self) -> None:
        self.declaracion = m15.leer_declaracion_tablas(
            _RAIZ_REPO / "config" / "informe_tablas.yaml")
        self.matrices = {c: e for c, e in self.declaracion.items()
                         if e.get("matriz")}

    def test_hay_matrices_declaradas(self) -> None:
        self.assertTrue(self.matrices)

    def test_toda_matriz_declara_sus_tres_campos_y_el_orden(self) -> None:
        for clave, entrada in self.matrices.items():
            matriz = entrada["matriz"]
            for campo in ("fila", "columna", "valor"):
                self.assertTrue(matriz.get(campo), f"{clave}: {campo}")
            self.assertTrue(matriz.get("orden"), clave)

    def test_el_orden_no_repite_valores(self) -> None:
        # Una columna repetida escribiria el mismo dato dos veces y dejaria
        # otra sin escribir.
        for clave, entrada in self.matrices.items():
            orden = [str(v) for v in entrada["matriz"]["orden"]]
            self.assertEqual(len(orden), len(set(orden)), clave)

    def test_una_matriz_no_declara_tambien_columnas(self) -> None:
        # Declarar las dos cosas deja en duda cual manda.
        for clave, entrada in self.matrices.items():
            self.assertFalse(entrada.get("columnas"), clave)


class PruebaLecturaDeTabla(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = Path(tempfile.mkdtemp())
        self.ruta = self.tmp / "datos.csv"
        self.ruta.write_text("a;b\n1;2\n3;4\n", encoding="utf-8-sig")

    def test_lee_las_filas(self) -> None:
        filas = m15.leer_tabla_csv(self.ruta, ";")
        self.assertEqual(len(filas), 2)
        self.assertEqual(filas[0]["a"], "1")

    def test_una_tabla_ausente_es_error(self) -> None:
        # Se detiene en lugar de dejar la tabla del estudio anterior en pie.
        with self.assertRaises(ErrorRutas):
            m15.leer_tabla_csv(self.tmp / "no_existe.csv", ";")



class PruebaEnsancharTabla(unittest.TestCase):
    """
    La plantilla viene dimensionada para cuatro microcuencas y el estudio tiene
    125. Crecer en filas y no en columnas dejaba media matriz fuera.
    """

    def _tabla(self, filas, columnas):
        import docx
        return docx.Document().add_table(rows=filas, cols=columnas)

    def test_ensancha_hasta_las_que_se_piden(self) -> None:
        tabla = self._tabla(3, 5)
        self.assertEqual(m15.anadir_columnas(tabla, 14), 9)
        self.assertEqual(len(tabla.columns), 14)

    def test_toda_fila_queda_con_el_mismo_ancho(self) -> None:
        # Una fila mas corta que las demas produce un documento que Word abre
        # con la tabla descuadrada.
        tabla = self._tabla(4, 5)
        m15.anadir_columnas(tabla, 14)
        for fila in tabla.rows:
            self.assertEqual(len(fila.cells), 14)

    def test_no_encoge_ni_toca_lo_que_ya_cabe(self) -> None:
        tabla = self._tabla(3, 9)
        self.assertEqual(m15.anadir_columnas(tabla, 5), 0)
        self.assertEqual(len(tabla.columns), 9)

    def test_la_cuadricula_crece_con_las_celdas(self) -> None:
        # Sin gridCol, Word reparte mal el ancho y la tabla se sale de la caja.
        from docx.oxml.ns import qn
        tabla = self._tabla(2, 5)
        m15.anadir_columnas(tabla, 14)
        cuadricula = tabla._tbl.find(qn("w:tblGrid"))
        self.assertEqual(len(cuadricula), 14)


class PruebaPlantillaConsolidada(unittest.TestCase):
    """
    LOS DEFECTOS CORREGIDOS NO PUEDEN VOLVER. Se arreglaron dentro del documento
    y un descuido al editarlo los devolveria en silencio: la tabla se llenaria
    igual y solo se veria al leer el informe terminado.
    """

    def setUp(self) -> None:
        self.plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not self.plantilla.is_file():
            self.skipTest("la plantilla no esta en el repositorio")
        import docx_plantilla as dp
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        self.documento = dp.abrir(self.plantilla)
        self.parrafos = []
        for hijo in self.documento.element.body.iterchildren():
            if hijo.tag.endswith("}p"):
                self.parrafos.append(Paragraph(hijo, self.documento))
            elif hijo.tag.endswith("}tbl"):
                for fila in Table(hijo, self.documento).rows:
                    for celda in fila.cells:
                        self.parrafos.extend(celda.paragraphs)

    def test_ninguna_figura_se_pide_con_un_nombre_que_la_cadena_no_usa(self):
        # La fase neutral del ENSO se llama 'neutral' en comun.oni; la
        # plantilla pedia 'neutro.png' y esa figura quedaba sin poner.
        for parrafo in self.parrafos:
            self.assertNotIn("neutro.png", parrafo.text)

    def test_cada_tabla_de_forma_titula_lo_que_muestra(self) -> None:
        # Tres tablas llevaban por encabezado 'Coeficiente de forma'; solo una
        # es de coeficiente de forma.
        titulos = [c.text.strip() for t in self.documento.tables
                   for c in t.rows[0].cells]
        self.assertEqual(titulos.count("Coeficiente de forma"), 1)

    def test_la_ventana_de_series_no_lleva_un_ano_fijo(self) -> None:
        # Decia 1980-2023, del informe de referencia, y la cadena calcula hasta
        # el ano del estudio.
        for parrafo in self.parrafos:
            self.assertNotIn("1980-2023", parrafo.text)

class PruebaNombreDeArchivoConTilde(unittest.TestCase):
    """
    El patron que lee el nombre de la figura tiene que admitir tildes.

    ERA UN TRUNCAMIENTO SILENCIOSO. Con la clase ASCII, una instruccion que
    pedia 'M14_comparacion_cambio_climatico.png' escrita con tildes se leia
    como 'tico.png', desde la ultima letra acentuada: el modulo buscaba un
    archivo con ESE nombre, no lo encontraba y reportaba una figura ausente, de
    modo que el diagnostico apuntaba al archivo y no a la instruccion. Aparecio
    de verdad al escribir la seccion de los dos escenarios de cambio climatico.
    """

    def test_lee_el_nombre_completo(self) -> None:
        self.assertEqual(
            m15.clasificar("Colocar Figura: M14_comparación_climático.png"),
            ("figura", "M14_comparación_climático.png"))

    def test_tambien_con_ene(self) -> None:
        self.assertEqual(
            m15.clasificar("Colocar Figura: M11_zonificación_año.png")[1],
            "M11_zonificación_año.png")

    def test_sigue_leyendo_los_nombres_sin_tilde(self) -> None:
        self.assertEqual(
            m15.clasificar("Colocar Figura: M14_qmax_vs_periodo.png")[1],
            "M14_qmax_vs_periodo.png")

    def test_no_se_lleva_el_texto_que_va_delante(self) -> None:
        # El patron busca dentro de lo que sigue a 'Colocar Figura:', y un
        # nombre de archivo no lleva espacios: si se los admitiera, arrastraria
        # la frase entera.
        self.assertEqual(
            m15.clasificar("Colocar Figura: el mapa M06_isoyetas.png")[1],
            "M06_isoyetas.png")


class PruebaSeccionDeVerificacion(unittest.TestCase):
    """
    La verificacion de crecientes va en 'Calibracion del Modelo'.

    NO ES UN DETALLE DE MAQUETACION. Lo que se presenta es lo que se hace
    cuando NO se puede calibrar, y colocado bajo 'Resultados' se leeria como un
    resultado del modelo en lugar de como su contraste. Paso de verdad: la
    primera insercion se ancló en una leyenda y el bloque quedó al otro lado
    del titulo, porque esa seccion no tiene ninguna tabla ni figura propia de
    la que colgarse.
    """

    def setUp(self) -> None:
        plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not plantilla.is_file():
            self.skipTest("la plantilla saneada no esta en el repositorio")
        import docx_plantilla as dp
        from docx.text.paragraph import Paragraph

        documento = dp.abrir(plantilla)
        self.parrafos = [Paragraph(hijo, documento)
                         for hijo in documento.element.body.iterchildren()
                         if hijo.tag.endswith("}p")]

    def _posicion(self, condicion) -> int:
        for indice, parrafo in enumerate(self.parrafos):
            if condicion(parrafo):
                return indice
        return -1

    def test_esta_antes_del_titulo_de_resultados(self) -> None:
        calibracion = self._posicion(
            lambda p: p.style.name == "Heading 3"
            and "Calibración del Modelo" in p.text)
        resultados = self._posicion(
            lambda p: p.style.name == "Heading 3"
            and p.text.strip() == "Resultados")
        tabla = self._posicion(
            lambda p: p.style.name == "Caption"
            and "Verificación de crecientes contra" in p.text)
        self.assertNotEqual(tabla, -1, "la leyenda no esta en la plantilla")
        self.assertLess(calibracion, tabla)
        self.assertLess(tabla, resultados)

    def test_las_tres_figuras_piden_archivos_distintos(self) -> None:
        # Es la misma clase de fallo que hubo en los dos escenarios de cambio
        # climatico: instrucciones distintas apuntando al mismo archivo.
        esperados = {"M14c_verificacion_J24.png", "M14c_verificacion_J29.png",
                     "M14c_media_movil_24h.png"}
        pedidos = {m15.clasificar(p.text)[1] for p in self.parrafos
                   if m15.clasificar(p.text)[0] == "figura"}
        self.assertTrue(esperados <= pedidos, esperados - pedidos)

    def test_explica_la_regulacion_aguas_arriba(self) -> None:
        """
        Sin esto, la tabla se leeria como que el modelo subestima.

        Las dos estaciones registran una corriente regulada por un embalse
        aguas arriba, de modo que su serie de medias diarias contabiliza
        descargas y vertimientos que el modelo de creciente no simula. El
        contraste acota, no valida, y el informe tiene que decirlo donde estan
        las cifras y no solo en una nota.
        """
        textos = " ".join(p.text for p in self.parrafos)
        self.assertIn("REGULACIÓN AGUAS ARRIBA", textos)
        self.assertIn("no permite concluir que el modelo subestime", textos)

    def test_dice_que_no_se_calibro(self) -> None:
        """
        Es la afirmacion que el informe tiene que sostener ante interventoria.

        Si los parametros se hubieran ajustado para reproducir estas cifras, la
        coincidencia dejaria de ser evidencia de que el modelo es adecuado.
        """
        textos = " ".join(p.text for p in self.parrafos)
        self.assertIn("no se realizó una calibración del modelo", textos)
        self.assertIn("sin modificar ningún parámetro", textos)


class PruebaParrafosDeTablas(unittest.TestCase):
    """
    Las comprobaciones tienen que mirar DENTRO de las tablas.

    ESA CEGUERA COSTO UN BLOQUE DUPLICADO EN EL ENTREGABLE. La plantilla trae
    las cuatro ilustraciones de isoyetas por fase en dos tablas de dos
    columnas, y el recorrido que se uso para buscarlas solo veia los parrafos
    de primer nivel: se dieron por inexistentes y se coloco otro bloque igual.
    La comprobacion de prosa heredada tenia el mismo punto ciego.
    """

    def _documento(self):
        import docx

        documento = docx.Document()
        documento.add_paragraph("en el cuerpo")
        tabla = documento.add_table(rows=1, cols=2)
        tabla.rows[0].cells[0].paragraphs[0].text = "dentro de una celda"
        tabla.rows[0].cells[1].add_paragraph("otra celda")
        return documento

    def test_recoge_los_de_las_celdas(self) -> None:
        textos = [p.text for p in
                  m15.parrafos_del_documento(self._documento())
                  if p.text.strip()]
        self.assertIn("en el cuerpo", textos)
        self.assertIn("dentro de una celda", textos)
        self.assertIn("otra celda", textos)

    def test_la_prosa_heredada_se_busca_tambien_en_las_tablas(self) -> None:
        documento = self._documento()
        documento.tables[0].rows[0].cells[0].paragraphs[0].text = (
            "La Quebrada No. 2 nace en el cerro.")
        hallados = m15.revisar_identidad(
            m15.parrafos_del_documento(documento),
            [{"termino": "Quebrada No. 2", "que_es": "la de referencia"}], [])
        self.assertEqual(len(hallados), 1)


class PruebaIdentidadAjena(unittest.TestCase):
    """
    La prosa heredada del informe de referencia.

    ES EL FALLO QUE NO SE VE. Un parrafo heredado es prosa correcta y bien
    escrita sobre OTRO proyecto: no es una instruccion en verde, ni una tabla
    sin llenar, ni una figura que falte. El informe de este estudio llego a
    decir que el IRH era 0,87 y el caudal ambiental 23,04 l/s, que son las
    cifras del estudio del que se copio la plantilla.
    """

    AJENOS = [
        {"termino": "Quebrada No. 2", "que_es": "la corriente de referencia"},
        {"termino": "Constructora Amarilo", "que_es": "el contratante"},
    ]

    class _Parrafo:
        def __init__(self, texto, estilo="Normal"):
            self.text = texto
            self.style = type("E", (), {"name": estilo})()

    def test_reporta_el_parrafo_con_el_termino_ajeno(self) -> None:
        parrafos = [self._Parrafo("La Quebrada No. 2 nace en el cerro.")]
        hallados = m15.revisar_identidad(parrafos, self.AJENOS, [])
        self.assertEqual(len(hallados), 1)
        self.assertEqual(hallados[0]["termino"], "Quebrada No. 2")

    def test_un_termino_declarado_como_propio_no_se_reporta(self) -> None:
        """
        Puede coincidir de verdad, y de hecho coincide.

        El contratante de este estudio es el mismo que el del de referencia.
        Senalarlo seria ruido, y un aviso que casi siempre sobra se aprende a
        ignorar: entonces deja de avisar de lo que importa.
        """
        parrafos = [self._Parrafo("Se recibió de Constructora Amarilo el dato.")]
        hallados = m15.revisar_identidad(parrafos, self.AJENOS,
                                         ["Constructora Amarilo"])
        self.assertEqual(hallados, [])

    def test_no_mira_las_leyendas_ni_las_instrucciones(self) -> None:
        # Una leyenda con el nombre ajeno se corrige al corregir su figura, y
        # el indice de tablas repite cada leyenda una segunda vez.
        parrafos = [
            self._Parrafo("Tabla 5-1. Quebrada No. 2", "Caption"),
            self._Parrafo("Colocar Figura: Quebrada No. 2.png"),
            self._Parrafo("Analizar Gráfico 5-1 de la Quebrada No. 2."),
        ]
        self.assertEqual(m15.revisar_identidad(parrafos, self.AJENOS, []), [])

    def test_un_parrafo_con_dos_terminos_se_reporta_dos_veces(self) -> None:
        # Cada uno exige su propia correccion; contarlo una vez escondería la
        # segunda.
        parrafos = [self._Parrafo("La Quebrada No. 2, de Constructora Amarilo.")]
        self.assertEqual(len(m15.revisar_identidad(parrafos, self.AJENOS, [])),
                         2)

    def test_sin_terminos_declarados_no_hay_nada_que_revisar(self) -> None:
        parrafos = [self._Parrafo("Un párrafo cualquiera.")]
        self.assertEqual(m15.revisar_identidad(parrafos, [], []), [])


class PruebaAnalisisRedactados(unittest.TestCase):
    """
    Los analisis del estudio sustituyen su instruccion en verde.

    VIVEN EN EL ESTUDIO Y NO EN LA PLANTILLA. Un analisis dice que significan
    los numeros de ESTE proyecto; escrito en la plantilla, que es compartida,
    los llevaria al siguiente estudio, que es justo el fallo que el informe ya
    tuvo con la prosa heredada.
    """

    def setUp(self) -> None:
        self.temporal = Path(tempfile.mkdtemp())

    def tearDown(self) -> None:
        import shutil
        shutil.rmtree(self.temporal, ignore_errors=True)

    def _archivo(self, contenido: str) -> Path:
        ruta = self.temporal / "analisis.yaml"
        ruta.write_text(contenido, encoding="utf-8")
        return ruta

    def test_se_indexa_por_la_instruccion_normalizada(self) -> None:
        # La instruccion lleva el numero que traia el informe del que se copio
        # el apartado, y es lo unico estable que hay para identificarla.
        ruta = self._archivo(
            "analisis:\n"
            "  - instruccion: \"Analizar Gráfico 5-1.\"\n"
            "    parrafos:\n"
            "      - \"El hidrograma crece de forma monótona.\"\n")
        leido = m15.leer_analisis(ruta)
        self.assertIn(m15._normalizar_leyenda("analizar grafico 5-1."), leido)

    def test_una_entrada_sin_parrafos_no_se_toma(self) -> None:
        # Dejaria la instruccion resuelta sin escribir nada, y el analisis
        # desapareceria del informe sin que nadie lo echara de menos.
        ruta = self._archivo(
            "analisis:\n"
            "  - instruccion: \"Analizar Gráfico 5-1.\"\n"
            "    parrafos: []\n")
        self.assertEqual(m15.leer_analisis(ruta), {})

    def test_un_archivo_ausente_no_es_error(self) -> None:
        # Significa que aun no se ha redactado ninguno, no que falte algo: las
        # instrucciones se quedan en verde y el modulo las cuenta.
        self.assertEqual(m15.leer_analisis(self.temporal / "no_existe.yaml"),
                         {})

    def test_la_instruccion_se_quita_y_no_se_reescribe(self) -> None:
        """
        Reutilizar el parrafo dejaria el analisis marcado en verde.

        El primer run de la instruccion lleva el resaltado, y escribir encima
        lo conserva: el consultor tendria que quitarlo a mano en cada uno y el
        recuento de pendientes seguiria contandolos.
        """
        import docx

        documento = docx.Document()
        documento.add_paragraph("antes")
        instruccion = documento.add_paragraph("Analizar Gráfico 5-1.")
        documento.add_paragraph("después")

        escritos = m15.resolver_analisis(
            instruccion, ["Primer párrafo.", "Segundo párrafo."], documento)

        self.assertEqual(escritos, 2)
        textos = [p.text for p in documento.paragraphs if p.text.strip()]
        self.assertEqual(textos,
                         ["antes", "Primer párrafo.", "Segundo párrafo.",
                          "después"])


class PruebaTandasDeGraficos(unittest.TestCase):
    """
    Una carpeta entera de gráficos, no una figura.

    La plantilla lo pide en ocho puntos y el contenido depende del estudio: en
    Refugio del Valle son 136 figuras, entre 8 y 32 por tanda. No puede
    escribirse en la plantilla, que no sabe cuántas estaciones tendrá el
    estudio siguiente.
    """

    def test_reconoce_la_instruccion_y_saca_la_subcarpeta(self) -> None:
        tipo, dato = m15.clasificar(
            "Agregar gráficos de la subcarpeta “precipitacion_cruda” de la "
            "carpeta de gráficos individuales de la carpeta de resultados y "
            "agregar un párrafo después de cada una.")
        self.assertEqual(tipo, "tanda")
        self.assertEqual(dato, "precipitacion_cruda")

    def test_una_instruccion_puede_pedir_dos_carpetas(self) -> None:
        # 'histograma_pdf' y 'papel de probabilidad' van en la misma.
        _, dato = m15.clasificar(
            "Agregar gráficos de la subcarpeta “histograma_pdf” y “papel de "
            "probabilidad” de la carpeta de gráficos individuales.")
        self.assertEqual(dato.split(","), ["histograma_pdf",
                                           "papel de probabilidad"])

    def test_la_instruccion_no_se_confunde_con_un_analisis(self) -> None:
        """
        'Agregar gráficos' pide figuras y 'Agregar análisis' pide texto.

        Contar la primera como redacción la daría por escrita cuando lo que
        falta es colocar 136 figuras.
        """
        self.assertEqual(
            m15.clasificar("Agregar análisis de Gráfico 4-3.")[0], "analisis")
        self.assertEqual(
            m15.clasificar("Agregar gráficos de la subcarpeta “enso” de la "
                           "carpeta de gráficos individuales.")[0], "tanda")

    def test_el_alias_resuelve_el_nombre_que_usa_la_instruccion(self) -> None:
        """
        La instrucción dice 'papel de probabilidad' y la carpeta se llama
        'papel_probabilidad'.

        Sin el alias la tanda no se colocaba y no había forma de saber por qué:
        solo faltaban dieciséis figuras.
        """
        tandas = m15.leer_tandas(
            _RAIZ_REPO / "config" / "informe_tandas.yaml")
        if not tandas:
            self.skipTest("no está la declaración de tandas")
        self.assertIn("papel_probabilidad", tandas)
        self.assertIn("papel_de_probabilidad", tandas)
        self.assertIs(tandas["papel_probabilidad"],
                      tandas["papel_de_probabilidad"])

    def test_cada_tanda_declara_leyenda_fuente_y_parrafo(self) -> None:
        # Sin cualquiera de los tres la figura entra suelta, sin poder
        # referenciarse o sin la explicación que la instrucción pide.
        tandas = m15.leer_tandas(
            _RAIZ_REPO / "config" / "informe_tandas.yaml")
        if not tandas:
            self.skipTest("no está la declaración de tandas")
        for clave, entrada in tandas.items():
            with self.subTest(tanda=clave):
                for campo in ("leyenda", "fuente", "parrafo", "nombra"):
                    self.assertTrue(str(entrada.get(campo, "")).strip(), campo)
                self.assertIn("{nombre}", str(entrada["leyenda"]))
                self.assertIn("{nombre}", str(entrada["parrafo"]))

    def test_el_nombre_sale_del_codigo_o_del_nombre_del_archivo(self) -> None:
        # Los módulos escriben unos por código y otros por nombre normalizado.
        estaciones = {"2120077": "TORCA", "apto_guaimaral_usta": "APTO GUAIMARAL"}
        self.assertEqual(
            m15.nombre_de_figura("2120077.png", "estacion", estaciones),
            "TORCA")
        self.assertEqual(
            m15.nombre_de_figura("precipitacion_total_historica_"
                                 "apto_guaimaral_usta.png", "estacion",
                                 estaciones),
            "APTO GUAIMARAL")

    def test_el_periodo_de_retorno_recupera_su_decimal(self) -> None:
        # El archivo no admite el punto y lo escribe como 'T2_33'.
        self.assertEqual(
            m15.nombre_de_figura("M11_mapa_pmax_T2_33.png", "periodo", {}),
            "2,33")
        self.assertEqual(
            m15.nombre_de_figura("M11_mapa_pmax_T100.png", "periodo", {}),
            "100")

    def test_un_archivo_que_no_se_reconoce_no_queda_sin_nombre(self) -> None:
        # Mejor el nombre del archivo que una leyenda vacía.
        self.assertEqual(
            m15.nombre_de_figura("otra_cosa.png", "estacion", {}), "otra cosa")


class PruebaCamposParaActualizar(unittest.TestCase):
    """
    Los campos del documento se marcan para que Word los recalcule.

    LAS LEYENDAS NO SE RENUMERAN SOLAS. Se componen con campos SEQ que
    conservan su ultimo resultado en cache, el del informe del que se derivo la
    plantilla: por eso unas dicen 'Tabla -.' y otras un numero que no les
    corresponde. Insertar una tabla nueva no cambia eso, y el consultor tenia
    que acordarse de pulsar Ctrl+E y F9.
    """

    def _documento(self):
        import docx

        documento = docx.Document()
        parrafo = documento.add_paragraph()
        for tipo in ("begin", "separate", "end"):
            run = parrafo.add_run()
            elemento = run._element.makeelement(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/"
                "main}fldChar", {})
            elemento.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/"
                "main}fldCharType", tipo)
            run._element.append(elemento)
        return documento

    def test_marca_el_campo_de_apertura(self) -> None:
        documento = self._documento()
        self.assertEqual(m15.marcar_campos_para_actualizar(documento), 1)

    def test_no_marca_el_separador_ni_el_cierre(self) -> None:
        # No llevan la instruccion del campo: marcarlos no significa nada y
        # ensucia el XML sin efecto.
        from docx.oxml.ns import qn

        documento = self._documento()
        m15.marcar_campos_para_actualizar(documento)
        marcados = [e.get(qn("w:fldCharType"))
                    for e in documento.element.body.iter()
                    if e.get(qn("w:dirty")) == "true"]
        self.assertEqual(marcados, ["begin"])

    def test_un_documento_sin_campos_no_es_error(self) -> None:
        import docx

        documento = docx.Document()
        documento.add_paragraph("sin campos")
        self.assertEqual(m15.marcar_campos_para_actualizar(documento), 0)


class PruebaFigurasDeEscenario(unittest.TestCase):
    """
    Las cinco figuras de los dos escenarios piden cinco archivos distintos.

    ES LA REGRESION QUE CASI SE ENTREGA. Al escribir la seccion, las cuatro
    instrucciones de figura apuntaban a DOS archivos: los dos hidrogramas al
    mismo y las dos curvas de Qmax al mismo. Los graficos de la pagina del
    escenario con factor habrian mostrado la curva del escenario sin el, y nada
    lo habria advertido, porque el archivo que nombraban existe: es el del otro
    escenario.
    """

    LEYENDAS = (
        "Hidrograma de Creciente Sitio de Proyecto (sin cambio climático)",
        "Qmax Vs. Periodo de Retorno (sin cambio climático)",
        "Hidrograma de Creciente Sitio de Proyecto (con cambio climático)",
        "Qmax Vs. Periodo de Retorno (con cambio climático)",
        "Comparación Caudales e Influencia",
    )

    def setUp(self) -> None:
        plantilla = _RAIZ_REPO / "templates" / "informe_base.docx"
        if not plantilla.is_file():
            self.skipTest("la plantilla saneada no esta en el repositorio")
        import docx_plantilla as dp
        from docx.text.paragraph import Paragraph

        documento = dp.abrir(plantilla)
        self.parrafos = [Paragraph(hijo, documento)
                         for hijo in documento.element.body.iterchildren()
                         if hijo.tag.endswith("}p")]

    def _archivo_de(self, leyenda: str):
        """El archivo que pide la instruccion que sigue a esa leyenda."""
        for indice, parrafo in enumerate(self.parrafos):
            if parrafo.style.name != "Caption" or leyenda not in parrafo.text:
                continue
            etiqueta = parrafo.text.strip().split()[0].rstrip(".")
            # LA MISMA FRASE ENCABEZA LA TABLA Y LA FIGURA del escenario, y la
            # tabla va primero: sin distinguir el prefijo se encuentra la de la
            # tabla, y detras de ella no hay instruccion de figura.
            if etiqueta not in ("Gráfico", "Ilustración", "Figura"):
                continue
            for siguiente in self.parrafos[indice + 1:indice + 3]:
                clase, dato = m15.clasificar(siguiente.text)[:2]
                if clase == "figura":
                    return dato
        return None

    def test_cada_leyenda_tiene_su_instruccion(self) -> None:
        for leyenda in self.LEYENDAS:
            with self.subTest(leyenda=leyenda):
                self.assertIsNotNone(self._archivo_de(leyenda))

    def test_las_cinco_piden_archivos_distintos(self) -> None:
        archivos = [self._archivo_de(leyenda) for leyenda in self.LEYENDAS]
        self.assertEqual(len(set(archivos)), len(self.LEYENDAS), archivos)

    def test_el_escenario_del_archivo_coincide_con_su_leyenda(self) -> None:
        """
        Comprobar que son distintos no basta: podrian estar intercambiados, y
        entonces cada pagina mostraria la figura de la otra. El sufijo
        '_referencia' es el del escenario sin factor.
        """
        for leyenda in self.LEYENDAS[:4]:
            with self.subTest(leyenda=leyenda):
                archivo = self._archivo_de(leyenda)
                sin_factor = "sin cambio climático" in leyenda
                self.assertEqual("_referencia" in archivo, sin_factor,
                                 f"{leyenda} -> {archivo}")


if __name__ == "__main__":
    unittest.main(verbosity=2)
