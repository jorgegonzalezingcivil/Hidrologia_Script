#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Aplica a la plantilla base las correcciones que hasta ahora se hacían al vuelo.

POR QUE EXISTE. Mientras la plantilla fue del consultor y no se podía tocar, sus
defectos se rodeaban desde config/informe_correcciones.yaml, que el M15 aplicaba
en cada corrida. Con el repositorio como fuente de verdad ese rodeo sobra: el
defecto se arregla en su sitio y la entrada declarada se retira.

NO ES PARTE DE LA CADENA. Se ejecuta a mano, una vez, y su efecto queda en el
historial de git, que es donde se puede revisar y deshacer. Volver a ejecutarlo
no hace daño: cada cambio se busca por su texto y si ya está aplicado no
encuentra nada.

    python tools/consolidar_plantilla.py            comprueba sin escribir
    python tools/consolidar_plantilla.py --escribir aplica y guarda
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

_RAIZ = Path(__file__).resolve().parents[1]
if str(_RAIZ / "src") not in sys.path:
    sys.path.insert(0, str(_RAIZ / "src"))

import docx_plantilla as plantilla_docx  # noqa: E402
import M15_informe as m15  # noqa: E402

PLANTILLA = _RAIZ / "templates" / "informe_base.docx"

# Nombres de figura que la plantilla pide con el nombre que la cadena no usa.
# Se buscan en el texto de la instruccion, sin tocar el resto del parrafo.
NOMBRES = (
    ("neutro.png", "neutral.png",
     "la cadena nombra la fase neutral del ENSO 'neutral' (comun.oni), no "
     "'neutro'"),
    ("Comparación Caudales e Influencia Cambo Climático",
     "Comparación Caudales e Influencia Cambio Climático",
     "decia 'Cambo Climatico'"),
)

# -----------------------------------------------------------------------------
# LAS FIGURAS DE LOS DOS ESCENARIOS DE CAMBIO CLIMATICO
#
# La seccion de resultados presenta los dos escenarios uno detras del otro, con
# tabla, hidrograma y curva de Qmax cada uno, y cierra con la figura de barras
# que los compara. Al escribirla, las CUATRO instrucciones de figura quedaron
# apuntando a DOS archivos: los dos hidrogramas al mismo y las dos curvas de
# Qmax al mismo. Los graficos de la pagina del escenario con factor habrian
# mostrado la curva del escenario sin el, y nada lo habria advertido: el nombre
# del archivo existe, solo que es el del otro escenario.
#
# Y el nombre que pedian no existia. La cadena escribe
# 'M14_hidrograma_Sink-1.png', donde el '1' es parte del nombre del elemento
# 'Sink-1' del modelo; al intercalar '_sin_cambio_climatico' antes del numero
# quedo partido en 'M14_hidrograma_Sink-_sin_cambio_climatico1.png'.
#
# Se identifican por la leyenda, que aqui SI las distingue una a una.
# -----------------------------------------------------------------------------
FIGURAS = (
    ("Hidrograma de Creciente Sitio de Proyecto (sin cambio climático)",
     "M14_hidrograma_Sink-1_referencia.png",
     "pedia el hidrograma del escenario con factor"),
    ("Qmax Vs. Periodo de Retorno (sin cambio climático)",
     "M14_qmax_vs_periodo_referencia.png",
     "pedia la curva del escenario con factor"),
    ("Hidrograma de Creciente Sitio de Proyecto (con cambio climático)",
     "M14_hidrograma_Sink-1.png",
     "el nombre del elemento del modelo es 'Sink-1' y quedo partido"),
    ("Qmax Vs. Periodo de Retorno (con cambio climático)",
     "M14_qmax_vs_periodo.png",
     "ya era el correcto; se declara para que las cuatro queden juntas"),
    ("Comparación Caudales e Influencia",
     "M14_escenarios_cc.png",
     "la cadena no pone tildes ni enes en los nombres de archivo"),
)


# Encabezados de tabla equivocados. Se identifican por la leyenda de SU tabla,
# porque el texto 'Coeficiente de forma' aparece bien en una y mal en dos.
ENCABEZADOS = (
    ("Coeficiente de compacidad microcuencas", "Coeficiente de forma",
     "Coeficiente de compacidad",
     "la tabla de compacidad titulaba su columna 'Coeficiente de forma'"),
    ("Índice de Sinuosidad microcuencas", "Coeficiente de forma",
     "Índice de sinuosidad",
     "la tabla de sinuosidad titulaba su columna 'Coeficiente de forma'"),
    ("Longitud Teórica de Series IDEAM",
     "Longitud Ventana de Tiempo 1980-2023 (años)",
     "Longitud Ventana de Tiempo desde 1980 (años)",
     "el literal 1980-2023 quedo del informe de referencia; la ventana llega "
     "hasta el ano del estudio y el parrafo que la introduce ya dice 'desde "
     "el ano 1980 hasta la fecha'"),
)


def _fijar_parrafo(parrafo, valor: str) -> None:
    """
    Reescribe un párrafo conservando el formato de su primer run.

    NO ES _fijar_texto DEL M15, que trabaja sobre CELDAS. Aquí hace falta la
    versión de párrafo, y conservar el primer run importa más que de costumbre:
    es el que lleva el resaltado verde con que la plantilla marca sus
    instrucciones, y perderlo dejaría la instrucción invisible para el
    consultor aunque el módulo la siguiera reconociendo por su texto.
    """
    if parrafo.runs:
        parrafo.runs[0].text = valor
        for sobrante in parrafo.runs[1:]:
            sobrante.text = ""
    else:
        parrafo.add_run(valor)


def _parrafos(documento):
    """Los del cuerpo y los de las celdas, en orden de documento."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    salida = []
    for hijo in documento.element.body.iterchildren():
        if hijo.tag.endswith("}p"):
            salida.append(Paragraph(hijo, documento))
        elif hijo.tag.endswith("}tbl"):
            for fila in Table(hijo, documento).rows:
                for celda in fila.cells:
                    salida.extend(celda.paragraphs)
    return salida


def _tablas_por_leyenda(documento):
    """Cada tabla instruida, indexada por la leyenda que la precede."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    elementos = []
    for hijo in documento.element.body.iterchildren():
        if hijo.tag.endswith("}p"):
            elementos.append(("p", Paragraph(hijo, documento)))
        elif hijo.tag.endswith("}tbl"):
            elementos.append(("t", Table(hijo, documento)))

    salida = {}
    for indice, (clase, objeto) in enumerate(elementos):
        if clase != "p" or m15.clasificar(objeto.text)[0] != "tabla":
            continue
        leyenda, tabla = "", None
        for siguiente in range(indice + 1, min(len(elementos), indice + 8)):
            if elementos[siguiente][0] == "t":
                tabla = elementos[siguiente][1]
                break
            texto = elementos[siguiente][1].text.strip()
            if texto and not leyenda:
                leyenda = texto
        if tabla is not None:
            salida[m15._normalizar_leyenda(leyenda)] = tabla
    return salida



# -----------------------------------------------------------------------------
# SECCION DE TIEMPO DE CONCENTRACION
#
# La plantilla documentaba nueve formulas y la cadena calcula trece. Medido
# contra data/referencia/tc_aplicabilidad.csv:
#
#   'SCS Ranser' y 'California Culverts Practice' son LA MISMA expresion,
#   0,947*(L^3/H)^0,385 con L en km frente a 0,0195*(L^3/H)^0,385 con L en m.
#   Las variables que la plantilla declara, L en km y H en m, lo confirman.
#   Decision del consultor: queda el nombre de la fuente con el suyo como alias.
#
#   'US Corps' estaba documentada y no se calcula. Decision del consultor:
#   se retira, porque documentar una formula que no entra en la mediana deja al
#   informe prometiendo algo que sus tablas no muestran.
#
#   Faltaban cinco por documentar. SCS Lag NO es SCS Ranser: es la ecuacion de
#   retardo del NRCS, la unica que usa el numero de curva.
# -----------------------------------------------------------------------------
QUITAR_FORMULAS = ("Fórmula de US Corps",)

RENOMBRAR = (
    ("Fórmula de SCS Ranser",
     "Fórmula de California Culverts Practice (SCS-Ranser)",
     "es la misma expresion que la cadena calcula como 'california'"),
)

# (titulo, expresion, lineas del 'Donde'). Todas devuelven horas.
FORMULAS_NUEVAS = (
    ("Fórmula de Johnstone y Cross",
     "tc = 0,4623 · L^0,5 · S^-0,25",
     ("tc = Tiempo de concentración (hr).",
      "L = Longitud del cauce principal (km).",
      "S = Pendiente media del cauce principal (m/m).")),
    ("Fórmula de Clark",
     "tc = 0,335 · (A / √S)^0,593",
     ("tc = Tiempo de concentración (hr).",
      "A = Área de la cuenca (km²).",
      "S = Pendiente media del cauce principal (m/m).")),
    ("Fórmula de Pilgrim y McDermott",
     "tc = 0,76 · A^0,38",
     ("tc = Tiempo de concentración (hr).",
      "A = Área de la cuenca (km²).")),
    ("Fórmula de Valencia y Zuluaga",
     "tc = 1,7694 · A^0,325 · L^-0,096 · S^-0,290",
     ("tc = Tiempo de concentración (hr).",
      "A = Área de la cuenca (km²).",
      "L = Longitud del cauce principal (km).",
      "S = Pendiente media del cauce principal (m/m).")),
    ("Fórmula de retardo del SCS (NRCS Lag)",
     "tc = tlag / 0,6,  con  tlag = (L^0,8 · (Sr + 1)^0,7) / (1900 · Y^0,5)",
     ("tc = Tiempo de concentración (hr).",
      "tlag = Tiempo de retardo (hr).",
      "L = Longitud del cauce principal (pies).",
      "Sr = Retención potencial máxima, Sr = 1000/CN - 10 (pulgadas).",
      "CN = Número de curva de escorrentía, adimensional.",
      "Y = Pendiente media de la cuenca (%).",
      "Es la única de la matriz que necesita el número de curva, y por eso es "
      "la coherente con el método de transformación del SCS adoptado en el "
      "modelo. Está definida para cuencas menores de 800 ha.")),
)


# LA PLANTILLA USA DOS ESTILOS PARA LO MISMO. Kirpich, Temez, Giandotti,
# California y Ventura Heras van en 'List Paragraph' y las demas en
# 'List Bullet'. Mirar uno solo hacia creer que la ultima formula era Passini, y
# las nuevas quedaban intercaladas en medio de la lista.
ESTILOS_DE_FORMULA = ("List Bullet", "List Paragraph")


def _es_titulo_de_formula(parrafo) -> bool:
    """Un encabezado de formula de la seccion, en cualquiera de sus estilos."""
    return (parrafo.style.name in ESTILOS_DE_FORMULA
            and parrafo.text.strip().lower().startswith("fórmula de"))


def _clonar(modelo, texto: str):
    """Copia un parrafo con su estilo y le pone otro texto."""
    import copy

    nuevo = copy.deepcopy(modelo._element)
    from docx.text.paragraph import Paragraph

    parrafo = Paragraph(nuevo, modelo._parent)
    _fijar_parrafo(parrafo, texto)
    return nuevo


def consolidar_tiempo_concentracion(documento, escribir: bool) -> list[str]:
    """Retira, renombra y completa las formulas de tiempo de concentracion."""
    from docx.text.paragraph import Paragraph

    hechos: list[str] = []
    parrafos = [Paragraph(h, documento)
                for h in documento.element.body.iterchildren()
                if h.tag.endswith("}p")]
    titulos = [i for i, p in enumerate(parrafos) if _es_titulo_de_formula(p)]
    if not titulos:
        return ["NO SE ENCONTRO ninguna formula en la plantilla"]

    # --- renombrar ---
    for viejo, nuevo, motivo in RENOMBRAR:
        for parrafo in parrafos:
            if parrafo.text.strip() != viejo:
                continue
            if escribir:
                _fijar_parrafo(parrafo, nuevo)
            hechos.append(f"'{viejo}' -> '{nuevo}' ({motivo})")

    # --- quitar el bloque entero, del titulo hasta el titulo siguiente ---
    for objetivo in QUITAR_FORMULAS:
        for orden, indice in enumerate(titulos):
            if parrafos[indice].text.strip() != objetivo:
                continue
            fin = (titulos[orden + 1] if orden + 1 < len(titulos)
                   else indice + 1)
            cuantos = fin - indice
            if escribir:
                for parrafo in parrafos[indice:fin]:
                    parrafo._element.getparent().remove(parrafo._element)
            hechos.append(f"retirada '{objetivo}' y su bloque "
                          f"({cuantos} parrafo(s))")

    # --- anadir las que faltan, detras de la ultima ---
    if escribir:
        parrafos = [Paragraph(h, documento)
                    for h in documento.element.body.iterchildren()
                    if h.tag.endswith("}p")]
        titulos = [i for i, p in enumerate(parrafos)
                   if _es_titulo_de_formula(p)]
    # EL BLOQUE TERMINA EN LA ULTIMA LINEA DE 'Donde', no en el primer parrafo
    # con otro estilo. Buscar el siguiente no-Normal saltaba por encima de la
    # nota en rosa y de la instruccion de la Tabla 3-18, y las formulas nuevas
    # se colaban ENTRE la instruccion y su tabla: el modulo dejaba de encontrarla
    # y la reportaba como 'sin tabla detras'.
    ultimo = titulos[-1]
    siguiente = ultimo + 1
    for indice in range(ultimo + 1, len(parrafos)):
        parrafo = parrafos[indice]
        if parrafo.style.name != "Normal":
            break
        if m15.clasificar(parrafo.text)[0]:
            break
        if any(r.font.highlight_color for r in parrafo.runs if r.text.strip()):
            break
        siguiente = indice + 1
    modelo_titulo = parrafos[ultimo]
    modelo_texto = parrafos[ultimo + 1] if ultimo + 1 < len(parrafos) else None

    presentes = {parrafos[i].text.strip() for i in titulos}
    ancla = parrafos[siguiente - 1]._element
    for titulo, expresion, donde in FORMULAS_NUEVAS:
        if titulo in presentes:
            continue
        if escribir and modelo_texto is not None:
            bloque = [_clonar(modelo_titulo, titulo),
                      _clonar(modelo_texto, expresion)]
            bloque += [_clonar(modelo_texto, "Donde:\t\t" + donde[0])]
            bloque += [_clonar(modelo_texto, linea) for linea in donde[1:]]
            for elemento in bloque:
                ancla.addnext(elemento)
                ancla = elemento
        hechos.append(f"anadida '{titulo}' con su expresion y sus variables")
    return hechos


# -----------------------------------------------------------------------------
# LAS CUATRO SECCIONES "REVISAR LA TEORIA... SEGUN LA METODOLOGIA ADOPTADA"
#
# Se contrasto cada parrafo contra el codigo que produce el resultado. Se
# encontraron tres desajustes reales, del mismo tipo que el de transito que
# senalo el consultor: la teoria describia un metodo que la cadena ya no usa
# asi, o callaba un metodo que si usa.
# -----------------------------------------------------------------------------
PARRAFOS_TEORIA = (
    # --- Infiltracion (Schosinsky y Losilla) --------------------------------
    # La plantilla decia que fc se tomaba de una tabla de Bradbury et al. 2000
    # (mm/hr, por textura). La cadena usa su propia tabla de doctrina
    # (data/referencia/infiltracion_schosinsky.csv), en mm/dia y por grupo
    # hidrologico SCS, declarada y sin validar. No es la misma fuente ni la
    # misma unidad.
    ("Se requiere de los valores de infiltración básica o capacidad de "
     "infiltración en el suelo, que al ser un parámetro especifico generado a "
     "partir de ensayos de laboratorio en diferentes puntos de la cuenca a "
     "través de estudios especializados es difícil de obtener, por lo cual se "
     "tomaron valores de referencia en la literatura y estudios en la región "
     "para suelos hidrológicos Tipo C.",
     "Se requiere el valor de infiltración básica del suelo (fc), que por ser "
     "un parámetro específico obtenido mediante ensayos de campo en distintos "
     "puntos de la cuenca resulta difícil de determinar de manera puntual. Se "
     "adopta un valor declarado por grupo hidrológico del suelo, según la "
     "clasificación del Servicio de Conservación de Suelos (SCS), expresado en "
     "milímetros por día. El grupo hidrológico de cada unidad de análisis se "
     "obtiene por muestreo del ráster de suelos sobre su área, y el valor de fc "
     "que le corresponde queda registrado en la tabla de doctrina del estudio, "
     "con su fuente y su condición de validación.",
     "la fuente y la unidad que la cadena usa (mm/dia, por grupo SCS) no son "
     "las de la tabla de Bradbury que sigue"),

    ("Teniendo en cuenta la ausencia de datos puntuales se asume como valor de "
     "infiltración básica el valor límite máximo de 38 mm/hr como condición "
     "crítica para los cálculos.",
     "La tabla anterior ilustra la relación general entre textura del suelo y "
     "capacidad de infiltración; los valores que este estudio adopta no "
     "provienen de ella, sino de la tabla de doctrina declarada por grupo "
     "hidrológico del suelo, expresada en milímetros por día. El grupo "
     "hidrológico dominante del área de estudio determina el valor de fc "
     "adoptado; cuando la cuenca presenta más de un grupo, la condición actual "
     "del cálculo toma uno solo para toda el área, y extenderlo por subcuenca "
     "es una decisión pendiente del consultor.",
     "la cadena no asume un valor fijo de 38 mm/hr: usa fc por grupo "
     "hidrologico, y hoy con UN grupo para toda la cuenca (ya anotado en "
     "informe_tablas.yaml). Sustituye la afirmacion sin dejar la tabla de "
     "Bradbury sin explicacion"),

    # --- Transito de crecientes ---------------------------------------------
    # Igual que en la tabla: el texto decia que el metodo usado era Muskingum,
    # y el que produce los caudales del modelo es Muskingum-Cunge. Y ataba
    # X = 0,2 "como el caso de estudio", cuando X se calcula por tramo.
    ("Existen varios métodos para propagación de crecientes en ríos y canales, "
     "pero el utilizado aquí es el método desarrollado por McCarthy en "
     "conexión con estudios de propagación del río Muskingum, en Ohio, Estados "
     "Unidos, este método hace uso de las ecuaciones de conservación de la "
     "masa.",
     "Existen varios métodos para propagación de crecientes en ríos y "
     "canales. Se calculan aquí dos: el método de Muskingum, desarrollado por "
     "McCarthy en conexión con estudios de propagación del río Muskingum, en "
     "Ohio, Estados Unidos, que hace uso de la ecuación de continuidad con "
     "parámetros K y X constantes; y el método de Muskingum-Cunge, que "
     "resuelve la misma ecuación pero deriva K y X de la hidráulica del tramo "
     "en cada subtramo y cada paso de tiempo, sin necesidad de calibrarlos "
     "contra hidrogramas observados. El método adoptado para los resultados "
     "del modelo es Muskingum-Cunge; Muskingum se presenta como "
     "parametrización alterna, de parámetros constantes.",
     "el modelo corre Muskingum-Cunge y el texto decia que el metodo usado "
     "aqui era Muskingum. Es la misma confusion que motivo la pregunta sobre "
     "la Tabla 5-10"),

    ("El parámetro k puede tomarse como el tiempo que le toma a la onda "
     "recorrer el total de la longitud de la sección en estudio. X es una "
     "constante que está entre 0 y 0.5, tomando como valor típico 0.2 para "
     "quebradas en zona relativamente planas como el caso de estudio.",
     "El parámetro K puede tomarse como el tiempo que le toma a la onda "
     "recorrer la longitud total del tramo en estudio. X es una constante "
     "entre 0 y 0,5 que refleja cuánto atenúa el tramo la onda de creciente: "
     "valores cercanos a 0,5 la trasladan casi sin atenuarla y valores "
     "cercanos a 0 la asimilan al comportamiento de un embalse de nivel "
     "horizontal. En ausencia de hidrogramas observados con que calibrarlos, K "
     "y X se obtienen aquí por linealización de Cunge en un caudal de "
     "referencia, a partir de la geometría, la pendiente y la resistencia "
     "hidráulica del tramo.",
     "ataba X = 0,2 'como el caso de estudio'; X se calcula por tramo y en "
     "este estudio la mediana da 0,4976, no 0,2"),
)

# Frases que se insertan despues de un parrafo ancla, porque describen algo
# que la cadena hace y la plantilla no menciona en absoluto.
INSERCIONES_TEORIA = (
    # ETP: Thornthwaite da el reparto mensual y Cenicafe, que solo depende de
    # la elevacion, da el nivel anual; la serie se ajusta a ese nivel. La
    # plantilla solo menciona Thornthwaite.
    ("Para la estación que cuenta con datos de temperatura mensual completados "
     "en la totalidad de meses del año se realizó el cálculo de este "
     "parámetro a nivel mensual el cual se presenta en el Anexo 4.",
     "Adicionalmente se calcula la evapotranspiración potencial anual "
     "mediante la ecuación de Cenicafé, que depende únicamente de la "
     "elevación y no requiere series de temperatura, y se emplea como "
     "estimador regional de contraste. La serie mensual obtenida por "
     "Thornthwaite se ajusta a ese nivel anual, de modo que conserva el "
     "reparto mensual que solo la temperatura puede aportar y corrige el "
     "nivel general, que Thornthwaite tiende a subestimar en climas fríos de "
     "montaña.",
     "la cadena ajusta Thornthwaite contra Cenicafe y la plantilla no lo "
     "menciona"),

    # ETR: se calcula tambien por Dekop, y la plantilla solo habla de Budyko.
    ("Ep = Evaporación potencial mensual (mm/mes).",
     "Se calcula también por el método de Dekop, versión simplificada de la "
     "misma familia de Budyko, que sirve de contraste: donde ambas "
     "formulaciones se separan, esa diferencia mide cuánto depende el "
     "resultado de la variante elegida. Budyko es el método adoptado para los "
     "resultados del balance; Dekop se presenta como verificación.",
     "la cadena calcula Budyko y Dekop y la plantilla solo describe Budyko"),
)


def consolidar_revision_teorica(documento, escribir: bool) -> list[str]:
    """Corrige y completa las cuatro secciones 'Revisar la teoría...'."""
    hechos: list[str] = []
    textos = {p.text.strip() for p in _parrafos(documento) if p.text.strip()}

    for viejo, nuevo, motivo in PARRAFOS_TEORIA:
        if nuevo in textos:
            hechos.append(f"YA APLICADO: {nuevo[:55]}...")
            continue
        parrafo = next((p for p in _parrafos(documento)
                        if p.text.strip() == viejo), None)
        if parrafo is None:
            hechos.append(f"NO SE ENCONTRO NI EL VIEJO NI EL NUEVO texto: "
                          f"{viejo[:55]}...")
            continue
        if escribir:
            _fijar_parrafo(parrafo, nuevo)
        hechos.append(f"parrafo corregido ({motivo})")

    for ancla, nuevo, motivo in INSERCIONES_TEORIA:
        if nuevo in textos:
            hechos.append(f"YA APLICADO: {nuevo[:55]}...")
            continue
        parrafo = next((p for p in _parrafos(documento)
                        if p.text.strip() == ancla), None)
        if parrafo is None:
            hechos.append(f"NO SE ENCONTRO el ancla: {ancla[:60]}...")
            continue
        if escribir:
            elemento = _clonar(parrafo, nuevo)
            parrafo._element.addnext(elemento)
        hechos.append(f"parrafo anadido tras '{ancla[:40]}...' ({motivo})")

    return hechos


# -----------------------------------------------------------------------------
# INSTRUCCIONES EN ROSA YA RESUELTAS: SE BORRAN DE LA PLANTILLA
#
# El rosa es distinto del verde. El verde marca donde va el dato de CADA
# estudio y tiene que sobrevivir para el siguiente; el rosa es una nota de
# autoria de la plantilla ("falta escribir esta teoria") y una vez escrita no
# tiene funcion futura, ni para este estudio ni para el proximo. Decision del
# consultor: se borra la instruccion, no solo su resaltado, porque dejar el
# texto sin marca lo confundiria con contenido del informe.
# -----------------------------------------------------------------------------
INSTRUCCIONES_ROSA_RESUELTAS = (
    "Revisar de la siguiente teoría cuáles autores nos falta por incluir "
    "según la programación hecha.",
    "Revisar la teoría de los siguientes párrafos y ajustar lo necesario "
    "según la metodología finalmente adoptada.",
    "Ajustar los parámetros de la Tabla 5-10 según la metodología adoptada.",
    "Incluir una tabla que relacione los autores con su viabilidad de uso en "
    "el cálculo según las características de las cuencas y el tipo de cuerpo "
    "de agua.",
)


def borrar_instrucciones_rosa_resueltas(documento, escribir: bool) -> list[str]:
    """Quita del documento las instrucciones rosa cuyo contenido ya se escribio."""
    hechos: list[str] = []
    for objetivo in INSTRUCCIONES_ROSA_RESUELTAS:
        coincidencias = [p for p in _parrafos(documento)
                         if p.text.strip() == objetivo]
        if not coincidencias:
            hechos.append(f"YA BORRADA: {objetivo[:55]}...")
            continue
        if escribir:
            for parrafo in coincidencias:
                parrafo._element.getparent().remove(parrafo._element)
        hechos.append(f"borrada(s) {len(coincidencias)} instruccion(es): "
                      f"{objetivo[:55]}...")
    return hechos


# -----------------------------------------------------------------------------
# TABLA DE VIABILIDAD DE AUTORES DE TIEMPO DE CONCENTRACION
#
# La pide la plantilla: "Incluir una tabla que relacione los autores con su
# viabilidad de uso en el calculo segun las caracteristicas de las cuencas y
# el tipo de cuerpo de agua." Es exactamente lo que ya vive en
# data/referencia/tc_aplicabilidad.csv: se lee de ahi y no se reinventa, para
# que la tabla del informe y la matriz que decide la mediana adoptada nunca
# puedan decir cosas distintas.
# -----------------------------------------------------------------------------
def consolidar_tabla_viabilidad(documento, escribir: bool) -> list[str]:
    """Inserta la tabla de autores de Tc con su rango de aplicabilidad."""
    import csv as _csv

    ancla_texto = ("Incluir una tabla que relacione los autores con su "
                   "viabilidad de uso en el cálculo según las características "
                   "de las cuencas y el tipo de cuerpo de agua.")
    ya_insertada = any(
        "Rango de área aplicable" in p.text for p in _parrafos(documento))
    if ya_insertada:
        return ["YA INSERTADA la tabla de viabilidad de autores de Tc"]

    ancla = next((p for p in _parrafos(documento)
                 if p.text.strip() == ancla_texto), None)
    if ancla is None:
        return ["NO SE ENCONTRO el ancla de la tabla de viabilidad de Tc"]

    ruta = _RAIZ / "data" / "referencia" / "tc_aplicabilidad.csv"
    filas = list(_csv.DictReader(ruta.open(encoding="utf-8-sig", newline=""),
                                 delimiter=";"))
    if not escribir:
        return [f"insertaria tabla de viabilidad con {len(filas)} autor(es)"]

    tabla = documento.add_table(rows=1 + len(filas), cols=3)
    tabla.style = "Grid Table 1 Light"
    encabezado = tabla.rows[0].cells
    for celda, titulo in zip(encabezado,
                             ("Autor", "Rango de área aplicable (km²)",
                              "Tipo de cuenca")):
        celda.text = titulo
    for indice, fila in enumerate(filas, start=1):
        celdas = tabla.rows[indice].cells
        desde = fila.get("area_min_km2", "").strip()
        hasta = fila.get("area_max_km2", "").strip()
        celdas[0].text = fila.get("nombre", "")
        celdas[1].text = f"{desde} a {hasta}" if desde or hasta else ""
        celdas[2].text = (fila.get("tipo_cuenca", "") or "").capitalize()

    ancla._element.addnext(tabla._tbl)
    return [f"tabla de viabilidad insertada con {len(filas)} autor(es), "
           f"leida de tc_aplicabilidad.csv"]


# -----------------------------------------------------------------------------
# LAS OCHO INSTRUCCIONES DE REDACCION NUEVA ("escribir", "incluir",
# "describir", "explicar", "agregar introduccion")
#
# Cada texto se verifico contra la funcion que produce el resultado
# (docstrings de M10, M05, M11/M12a/M12b, M18a) y queda en teoria neutra, sin
# cifras de este estudio: es reutilizable para el siguiente.
#
# La instruccion se BORRA y el contenido se inserta como parrafo NUEVO (no se
# reescribe el parrafo de la instruccion), para que el texto final no herede
# el resaltado rosa de su run.
# -----------------------------------------------------------------------------
REDACCION_NUEVA = (
    ("Describir de manera general de donde se ha tomado la información del "
     "tipo de suelo hidrológico.",
     ("El grupo hidrológico del suelo, que el método del número de curva "
      "exige, se obtiene del producto HYSOGs250m (Ross et al., 2018, Global "
      "Hydrologic Soil Groups, ORNL DAAC), un ráster global a escala "
      "1:250.000 que asigna directamente el grupo del Servicio de "
      "Conservación de Suelos (A, B, C o D) a cada celda, incluidas las "
      "combinaciones duales que dependen de si el suelo está drenado. Es un "
      "producto de origen global y no un levantamiento de campo, de modo que "
      "su resolución condiciona la escala a la que puede leerse el resultado "
      "dentro de la cuenca.",),
     "fuente y escala del rater de suelos (config.yaml, referencia_nacional)"),

    ("Escribir la teoría de la curva hipsométrica y para qué sirve.",
     ("La curva hipsométrica representa la fracción del área de la cuenca "
      "que se encuentra por encima de cada cota, normalizada entre 0 y 1 en "
      "ambos ejes, en la forma que la definió Strahler. Su forma describe la "
      "etapa erosiva de la cuenca: una curva convexa, con la mayor parte del "
      "área en las cotas altas, corresponde a una cuenca joven en "
      "desequilibrio; una curva cóncava, con la mayor parte del área en las "
      "cotas bajas, corresponde a una cuenca erosionada; y una forma "
      "intermedia, en S, corresponde a una cuenca madura. La integral de la "
      "curva resume esa lectura en un solo número: valores por encima de "
      "0,60 se asocian a cuencas jóvenes, entre 0,35 y 0,60 a cuencas "
      "maduras, y por debajo de 0,35 a cuencas erosionadas.",),
     "docstring de curva_hipsometrica() en M10_morfometria.py"),

    ("Incluir texto teórico del análisis de dobles masas.",
     ("El análisis de dobles masas compara la serie acumulada de cada "
      "estación contra la serie acumulada de un patrón construido con las "
      "estaciones vecinas, en el mismo periodo. Mientras la relación entre "
      "las dos series se mantenga estable, la curva de dobles masas es una "
      "línea recta; un cambio de pendiente indica que la estación empezó a "
      "registrar en una escala distinta de la del patrón, por un cambio de "
      "instrumento, de ubicación o de operador. El tramo anterior al quiebre "
      "se corrige multiplicándolo por la razón entre las dos pendientes, de "
      "modo que quede en la misma escala que el tramo reciente, que es el "
      "que refleja las condiciones actuales de la estación.",),
     "docstring de corregir_por_doble_masa() en M05_precipitacion_mensual.py"),

    ("Incluir texto teórico del análisis de datos anómalos.",
     ("La detección de datos anómalos se hace mes a mes y no sobre la serie "
      "completa: en un régimen de lluvias bimodal, un solo rango de "
      "referencia para los doce meses marcaría como anómala toda una "
      "temporada húmeda, porque la comparación tiene que hacerse contra el "
      "mismo mes calendario en los demás años. El método adoptado es el "
      "rango intercuartílico (IQR): se calculan el primer y el tercer "
      "cuartil de los valores históricos de cada mes y se marcan como "
      "anómalos los que caen fuera de los límites que resultan de extender "
      "ese rango un múltiplo declarado. Un dato marcado no se descarta "
      "automáticamente: queda señalado para que el consultor decida su "
      "tratamiento.",),
     "detectar_anomalos_por_mes() y config.anomalos en M05_precipitacion_mensual.py"),

    ("Incluir texto teórico del análisis de validación estadística realizado "
     "en el complemento de datos.",
     ("El complemento de los datos faltantes se hace por regresión lineal "
      "simple contra la estación vecina mejor correlacionada que tenga dato "
      "disponible en ese periodo, de modo que cada valor completado se "
      "apoye en la mejor información disponible y no en un promedio que "
      "diluya la relación entre estaciones. La validez del método se "
      "comprueba por validación cruzada: se retira de la serie una muestra "
      "de datos conocidos, se estiman con el mismo procedimiento de "
      "complemento y se compara el valor estimado contra el valor real "
      "retirado, lo que permite cuantificar el error del complemento sobre "
      "datos cuyo valor verdadero sí se conoce.",),
     "_regresion_vecinas() y validacion_cruzada() en M05_precipitacion_mensual.py"),

    ("Explicar por qué se hizo la zonificación de pluviómetros para HEC HMS "
     "y cuál es el procedimiento y paso de la precipitación de 24 horas a 3 "
     "horas.",
     ("La zonificación agrupa las subcuencas en un número reducido de zonas "
      "pluviométricas, cada una representada por un único hietograma. Es "
      "necesaria porque en HEC-HMS cada hietograma distinto exige su propio "
      "pluviómetro: asignar una serie de diseño distinta a cada subcuenca y "
      "cada periodo de retorno multiplicaría el número de series hasta un "
      "punto que ningún consultor puede mantener ni revisar, mientras que "
      "agrupar subcuencas cuya lámina de diseño no difiere más de un umbral "
      "declarado reduce esa cantidad a un número manejable sin perder la "
      "variación espacial de la lluvia dentro de la cuenca. La lámina de "
      "cada zona se calcula como el promedio ponderado por área de las "
      "subcuencas que la componen, para que una subcuenca grande no pese lo "
      "mismo que una pequeña.",
      "El paso de la precipitación máxima en 24 horas a la duración de "
      "diseño de 3 horas se resuelve con tres hipótesis, calculadas en "
      "paralelo y sin que ninguna se adopte por defecto: la primera asigna "
      "la lámina completa de 24 horas a la duración de diseño, y es la más "
      "conservadora; la segunda integra la curva IDF adoptada sobre esa "
      "duración; y la tercera aplica un factor de escala declarado sobre la "
      "lámina de 24 horas. Las tres se calculan porque la diferencia entre "
      "ellas es la que permite decidir con criterio cuál adoptar, y esa "
      "decisión debe quedar registrada."),
     "agrupar_por_zona() en M12b_hietogramas.py y desagregar() en M12a_idf.py"),

    ("Escribir procedimiento para procesar la información de temperatura de "
     "las estaciones.",
     ("El procesamiento de la información de temperatura parte de las "
      "estaciones con registros de temperatura máxima y mínima diaria "
      "dentro del área de influencia. Con esos datos se calcula la "
      "temperatura media mensual de cada estación y se ajusta, por mínimos "
      "cuadrados, un gradiente altitudinal propio del estudio, en la forma "
      "T = a + b·h, con h la elevación de la estación. El ajuste se reporta "
      "con el intervalo de confianza de su pendiente y no solo con el "
      "coeficiente de determinación, porque con pocas estaciones y un rango "
      "de elevación estrecho el R² puede salir alto sin que la pendiente "
      "esté bien determinada, y es la pendiente la que se extrapola sobre "
      "las partes altas de la cuenca donde no hay estación. El gradiente se "
      "contrasta contra valores de referencia de la región antes de "
      "adoptarse.",),
     "ajustar_gradiente() en M18a_temperatura.py"),

    ("Escribir la teoría de cómo se hace la calibración, variando los "
     "números de curva, tránsito de Muskingum-cunge y tiempo de rezago.",
     ("La calibración de un modelo lluvia-escorrentía consiste en ajustar, "
      "dentro de su rango físicamente justificado, los parámetros que el "
      "modelo no puede medir directamente (el número de curva, los "
      "parámetros del tránsito de Muskingum-Cunge y el tiempo de rezago) "
      "hasta que el hidrograma simulado reproduzca el hidrograma observado "
      "en un punto de control. La comparación se hace con métricas "
      "objetivas, como el coeficiente de Nash-Sutcliffe, el error "
      "cuadrático medio y el sesgo porcentual (PBIAS), y no solo por "
      "inspección visual. La calibración exige series de caudal observadas "
      "en el punto de control o cerca de él; sin series limnigráficas o "
      "limnimétricas utilizables, el modelo se deja sin calibrar y los "
      "parámetros quedan en los valores que la caracterización morfométrica "
      "y la doctrina adoptada determinan, lo que debe declararse "
      "explícitamente en el informe.",),
     "config.calibracion (metricas) y CLAUDE.md seccion 6"),

    ("Agregar Introducción al numeral.",
     ("Se presentan a continuación los resultados del balance hídrico a "
      "largo plazo, obtenidos de aplicar el método adoptado a la "
      "información hidroclimatológica caracterizada en los numerales "
      "anteriores.",),
     "introduccion generica del numeral Resultados del balance"),
)


# -----------------------------------------------------------------------------
# EL ANALISIS DE LOS DOS ESCENARIOS DE CAMBIO CLIMATICO
#
# Va con la misma regla que las ocho anteriores: TEORIA NEUTRA, sin cifras de
# este estudio. Las cifras las ponen las dos tablas y la figura de barras, y el
# texto tiene que seguir sirviendo cuando el factor sea otro. Lo que si dice, y
# es lo que no se puede deducir mirando las tablas, es POR QUE el caudal no
# crece en la misma proporcion que la lluvia y CUAL de los dos escenarios es el
# de diseno.
# -----------------------------------------------------------------------------
ANALISIS_ESCENARIOS = (
    "Escribir análisis sobre resultados sin cambio climático y cambio "
    "climático",
    ("Los resultados se presentan en dos escenarios. El escenario de "
     "referencia se obtiene con la lluvia de diseño derivada directamente de "
     "los registros históricos de las estaciones, sin ninguna corrección; el "
     "escenario de diseño aplica sobre esa misma lluvia el factor de cambio "
     "climático que corresponde al periodo de retorno y a la ventana de "
     "proyección adoptada. Los dos se calculan con el mismo modelo, la misma "
     "distribución temporal de la tormenta, los mismos números de curva y el "
     "mismo tránsito: la única diferencia entre ellos es la lámina de "
     "precipitación.",
     "El incremento del caudal no es proporcional al incremento de la lluvia, "
     "y esa es la razón de calcular el segundo escenario en lugar de escalar "
     "el primero. El método del número de curva descuenta una abstracción "
     "inicial antes de que se produzca escorrentía, de modo que la fracción de "
     "la lluvia que se convierte en caudal aumenta con la magnitud del evento: "
     "un mismo porcentaje adicional de precipitación produce un porcentaje "
     "mayor de caudal, y ese efecto es más acentuado en las crecientes "
     "frecuentes, donde la lluvia total es pequeña frente al umbral de "
     "pérdidas, que en las extraordinarias. En consecuencia, el aporte "
     "relativo del factor de cambio climático disminuye a medida que crece el "
     "periodo de retorno, aunque su aporte absoluto aumente.",
     "El escenario de diseño, con el factor aplicado, es el que se adopta "
     "para el dimensionamiento hidráulico de las obras, en aplicación de la "
     "regla condicional que incorpora la proyección climática únicamente "
     "cuando es de incremento. El escenario de referencia no constituye una "
     "alternativa de diseño entre las que elegir: se presenta para separar "
     "explícitamente la parte del caudal que proviene del registro observado "
     "de la parte que proviene de la proyección, y así permitir que una "
     "revisión posterior evalúe por separado el modelo hidrológico y la "
     "hipótesis climática que se le superpone."),
    "el consultor pidio el analisis de los dos escenarios",
)


def consolidar_redaccion_nueva(documento, escribir: bool) -> list[str]:
    """Redacta las ocho secciones y borra su instruccion, en una sola pasada."""
    hechos: list[str] = []
    textos_actuales = {p.text.strip() for p in _parrafos(documento)
                       if p.text.strip()}

    for ancla_texto, parrafos_nuevos, motivo in (
            REDACCION_NUEVA + (ANALISIS_ESCENARIOS,)):
        if parrafos_nuevos[0] in textos_actuales:
            hechos.append(f"YA REDACTADO: {ancla_texto[:50]}...")
            continue
        ancla = next((p for p in _parrafos(documento)
                     if p.text.strip() == ancla_texto), None)
        if ancla is None:
            hechos.append(f"NO SE ENCONTRO el ancla: {ancla_texto[:50]}...")
            continue
        if escribir:
            punto = ancla._element
            for texto in parrafos_nuevos:
                nuevo = documento.add_paragraph(texto, style="Normal")
                punto.addnext(nuevo._element)
                punto = nuevo._element
            ancla._element.getparent().remove(ancla._element)
        hechos.append(f"redactada ({motivo})")
    return hechos

def consolidar(ruta: Path, escribir: bool) -> list[str]:
    """Aplica los cambios y devuelve lo que hizo, o lo que haría."""
    documento = plantilla_docx.abrir(ruta)
    hechos: list[str] = []

    # --- 1. Figuras que la plantilla pide mal --------------------------------
    # SE REUSA EL MISMO EMPAREJADOR QUE EL M15. La leyenda sola no identifica la
    # instruccion, porque 'Areas microcuencas' encabeza tres; el desempate por
    # el archivo que hoy nombra ya esta resuelto en planear_correcciones.
    correcciones = m15.leer_correcciones(
        _RAIZ / "config" / "informe_correcciones.yaml")
    plan, ambiguas, sin_uso = m15.planear_correcciones(documento, correcciones)
    if ambiguas:
        raise SystemExit(f"correcciones ambiguas, no se toca nada: {ambiguas}")
    for parrafo in _parrafos(documento):
        entrada = plan.get(parrafo._element)
        if entrada is None:
            continue
        pedia = m15.clasificar(parrafo.text)[1]
        nuevo = str(entrada["archivo"])
        if escribir:
            _fijar_parrafo(parrafo, parrafo.text.replace(pedia, nuevo))
        hechos.append(f"figura bajo '{entrada.get('leyenda')}': "
                      f"{pedia} -> {nuevo}")
    for clave in sin_uso:
        hechos.append(f"YA APLICADA, retirar de informe_correcciones.yaml: "
                      f"{clave}")

    # --- 2. Nombres de archivo que la cadena no usa --------------------------
    for viejo, nuevo, motivo in NOMBRES:
        for parrafo in _parrafos(documento):
            if viejo not in parrafo.text:
                continue
            if escribir:
                _fijar_parrafo(parrafo, parrafo.text.replace(viejo, nuevo))
            hechos.append(f"{viejo} -> {nuevo} ({motivo})")

    # --- 3. Encabezados de tabla equivocados ---------------------------------
    tablas = _tablas_por_leyenda(documento)
    for leyenda, viejo, nuevo, motivo in ENCABEZADOS:
        tabla = tablas.get(m15._normalizar_leyenda(leyenda))
        if tabla is None:
            hechos.append(f"NO SE ENCONTRO la tabla '{leyenda}'")
            continue
        for celda in tabla.rows[0].cells:
            if celda.text.strip() != viejo:
                continue
            if escribir:
                m15._fijar_texto(celda, nuevo)
            hechos.append(f"'{leyenda}': encabezado '{viejo}' -> '{nuevo}' "
                          f"({motivo})")

    # --- 4. Seccion de tiempo de concentracion -------------------------------
    hechos.extend(consolidar_tiempo_concentracion(documento, escribir))

    # --- 5. Las cuatro secciones de revision teorica -------------------------
    hechos.extend(consolidar_revision_teorica(documento, escribir))

    # --- 6. Tabla de viabilidad de autores de Tc -----------------------------
    hechos.extend(consolidar_tabla_viabilidad(documento, escribir))

    # --- 6b. Las ocho instrucciones de redaccion nueva -----------------------
    hechos.extend(consolidar_redaccion_nueva(documento, escribir))

    # --- 6d. Figuras de los dos escenarios de cambio climatico ---------------
    hechos.extend(consolidar_figuras_de_escenario(documento, escribir))

    # --- 6c. Tablas que la cadena produce y la plantilla no declaraba --------
    hechos.extend(consolidar_bloques_nuevos(documento, escribir))

    # --- 7. Instrucciones rosa ya resueltas: se borran -----------------------
    # VA AL FINAL: tiene que ver escritas las secciones 1 a 6 antes de borrar
    # sus notas, o borraria la instruccion sin haber comprobado que el
    # contenido que la reemplaza ya esta.
    hechos.extend(borrar_instrucciones_rosa_resueltas(documento, escribir))

    if escribir and hechos:
        documento.save(str(ruta))
    return hechos


def _indexar(documento):
    """(clase, objeto) de cada hijo del cuerpo, en orden de documento."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    salida = []
    for hijo in documento.element.body.iterchildren():
        if hijo.tag.endswith("}p"):
            salida.append(("p", Paragraph(hijo, documento)))
        elif hijo.tag.endswith("}tbl"):
            salida.append(("t", Table(hijo, documento)))
    return salida


def _leyenda(objetos, texto: str):
    """
    La leyenda de estilo 'Caption' que dice ese texto, y su posicion.

    SE EXIGE EL ESTILO. El indice de tablas repite cada leyenda con estilo
    'table of figures', de modo que buscar solo por texto encuentra dos y la
    primera es la del indice, al principio del documento.
    """
    for indice, (clase, objeto) in enumerate(objetos):
        if (clase == "p" and objeto.style.name == "Caption"
                and texto in objeto.text):
            return indice, objeto
    return -1, None


def _clonar_leyenda(modelo, leyenda: str):
    """
    Copia una leyenda conservando sus campos de numeracion.

    NO SIRVE _clonar. Esa version escribe en el primer run y vacia los demas,
    y en una leyenda los runs del medio son los campos SEQ que producen el
    numero: vaciarlos deja el campo roto y la leyenda sin numerar. Aqui se
    conserva todo hasta el ultimo campo y se sustituye solo el texto que sigue,
    con lo que la leyenda nueva se numera sola y entra en la secuencia del
    capitulo, sin que haya que renumerar las que ya estaban.
    """
    import copy

    from docx.text.paragraph import Paragraph

    parrafo = Paragraph(copy.deepcopy(modelo._element), modelo._parent)
    puesto = False
    for run in parrafo.runs:
        xml = run._element.xml
        if "fldChar" in xml or "instrText" in xml:
            continue
        if not run.text.strip().startswith("."):
            continue
        run.text = f". {leyenda}"
        puesto = True
        break
    if not puesto:
        raise SystemExit(f"la leyenda modelo no tiene texto tras su numero: "
                         f"{modelo.text!r}")
    # Lo que venia detras del texto sustituido era el resto de la leyenda
    # vieja, partido en varios runs por el corrector ortografico de Word.
    visto = False
    for run in parrafo.runs:
        if run.text == f". {leyenda}":
            visto = True
            continue
        if visto and "fldChar" not in run._element.xml:
            run.text = ""
    return parrafo._element


def _clonar_tabla(modelo, columnas: int):
    """
    Copia una tabla de la plantilla dejandola con otro numero de columnas.

    SE CLONA Y NO SE CREA. Una tabla nueva con 'add_table' sale con el estilo
    'Normal Table', que en esta plantilla no lleva bordes ni sombreado de
    encabezado: el formato de casa esta aplicado celda por celda, no en un
    estilo con nombre, de modo que la unica forma de heredarlo es copiar una
    tabla que ya lo tenga.

    Se quitan columnas por la derecha y se reparte el ancho entre las que
    quedan, para que la tabla siga ocupando el ancho util de la pagina.
    """
    import copy

    from docx.oxml.ns import qn
    from docx.table import Table

    if columnas > len(modelo.columns):
        raise SystemExit(f"la tabla modelo tiene {len(modelo.columns)} "
                         f"columna(s) y se piden {columnas}")
    tabla = Table(copy.deepcopy(modelo._tbl), modelo._parent)
    rejilla = tabla._tbl.find(qn("w:tblGrid"))
    definiciones = rejilla.findall(qn("w:gridCol"))
    total = sum(int(d.get(qn("w:w")) or 0) for d in definiciones)
    for sobrante in definiciones[columnas:]:
        rejilla.remove(sobrante)
    for definicion in rejilla.findall(qn("w:gridCol")):
        definicion.set(qn("w:w"), str(total // columnas))
    for fila in tabla._tbl.findall(qn("w:tr")):
        celdas = fila.findall(qn("w:tc"))
        for sobrante in celdas[columnas:]:
            fila.remove(sobrante)
        for celda in fila.findall(qn("w:tc")):
            propiedades = celda.find(qn("w:tcPr"))
            ancho = (propiedades.find(qn("w:tcW"))
                     if propiedades is not None else None)
            if ancho is not None:
                # El ancho de celda va en porcentaje del ancho de la tabla,
                # que es 5000 = 100 %. El de la rejilla va en twips.
                ancho.set(qn("w:w"), str(5000 // columnas))
    return tabla


# -----------------------------------------------------------------------------
# TABLAS QUE LA CADENA YA PRODUCE Y LA PLANTILLA NO DECLARA
#
# El M15 no inventa estructura: llena lo que la plantilla pide. Una tabla
# declarada en config/informe_tablas.yaml cuya leyenda no exista en el
# documento NO produce error al generar el informe, simplemente no aparece, y
# eso no se nota al revisarlo. Tres pruebas de tests/test_m15.py lo impiden, y
# estaban en rojo por estas dos.
#
# LA LEYENDA Y LOS ENCABEZADOS SE LEEN DE informe_tablas.yaml, no se copian
# aqui. Es la misma razon por la que la tabla de viabilidad se lee de
# tc_aplicabilidad.csv: si se escribieran dos veces, el documento y la
# declaracion podrian acabar diciendo cosas distintas, y el emparejamiento del
# M15 es por leyenda exacta.
#
# EL NUMERO DE LA INSTRUCCION ES INFORMATIVO. La leyenda se numera con campos
# SEQ y Word calcula el numero al actualizarlos; el M15 empareja por el texto,
# nunca por el numero (ver M15_informe.leer_declaracion_tablas). Los que van
# aqui son los que quedan al insertar: el capitulo 5 llegaba a la 5-12, el
# embalse entra antes en 'Construccion del Modelo' y corre las siguientes.
# -----------------------------------------------------------------------------
BLOQUES_NUEVOS = (
    {
        "leyenda": "Curva de almacenamiento y descarga del embalse",
        # Entra detras de la 5-10, 'Datos transito crecientes microcuencas',
        # que es la ultima de 'Construccion del Modelo'. Con eso las dos de
        # resultados, que el consultor rotulo 5-11 y 5-12, pasan a ser la 5-12
        # y la 5-13 en cuanto Word actualice los campos.
        "numero": "5-11",
        # Va donde se arma el modelo y no donde se presentan sus resultados: el
        # embalse es un elemento del .basin, y el estado de operacion del que
        # parte es la decision con mas peso de todo el transito.
        "antes_de": "Modelo HEC-HMS",
        "fuente": "Fuente: EAAB, INCOHISA, 2024.",
        "motivo": "la curva del embalse no estaba declarada en la plantilla",
    },
)

INSTRUCCION = ("Completar la Tabla {numero} con datos del estudio, "
               "conservando el estilo de letra y tamaño, cambiar a color "
               "negro al finalizar. Agregar las columnas o filas que sean "
               "necesarias.")


def _leyenda_de_figura(objetos, texto: str):
    """
    La leyenda de FIGURA que dice ese texto, y su posicion.

    NO SIRVE _leyenda. La misma frase encabeza la tabla y la figura del mismo
    escenario, 'Tabla -. Qmax Vs. Periodo de Retorno (sin cambio climatico)' y
    'Grafico -. Qmax ...', y la tabla va primero: buscando por texto se
    encontraba la de la tabla y detras habia una tabla, no una instruccion de
    figura.
    """
    for indice, (clase, objeto) in enumerate(objetos):
        if clase != "p" or objeto.style.name != "Caption":
            continue
        etiqueta = objeto.text.strip().split()[0] if objeto.text.strip() else ""
        if etiqueta.rstrip(".") not in ("Gráfico", "Ilustración", "Figura"):
            continue
        if texto in objeto.text:
            return indice, objeto
    return -1, None


def consolidar_figuras_de_escenario(documento, escribir: bool) -> list[str]:
    """
    Pone en cada instruccion de figura el archivo que le corresponde.

    EN LAS FIGURAS LA LEYENDA VA ANTES de la instruccion, al contrario que en
    las tablas. No es un descuido de la plantilla: es su composicion.

    SE REESCRIBE LA INSTRUCCION ENTERA y no se sustituye el nombre dentro del
    texto. El nombre que la plantilla traia lleva tildes, y sustituir el trozo
    que el patron del M15 alcanza a leer dejaria pegado lo que quedaba delante.
    """
    hechos: list[str] = []
    objetos = _indexar(documento)
    for leyenda, archivo, motivo in FIGURAS:
        indice, _ = _leyenda_de_figura(objetos, leyenda)
        if indice < 0:
            hechos.append(f"NO SE ENCONTRO la leyenda de figura '{leyenda}'")
            continue
        instruccion = None
        for siguiente in range(indice + 1, min(len(objetos), indice + 4)):
            clase, objeto = objetos[siguiente]
            if clase != "p":
                break
            if m15.clasificar(objeto.text)[0] == "figura":
                instruccion = objeto
                break
        if instruccion is None:
            hechos.append(f"'{leyenda}' no tiene instruccion de figura detras")
            continue
        pedia = instruccion.text.strip()
        quiere = f"Colocar Figura: {archivo}"
        if pedia == quiere:
            hechos.append(f"YA CORREGIDA: '{leyenda}' pide {archivo}")
            continue
        if escribir:
            _fijar_parrafo(instruccion, quiere)
        hechos.append(f"'{leyenda}': {pedia} -> {quiere} ({motivo})")
    return hechos


def consolidar_bloques_nuevos(documento, escribir: bool) -> list[str]:
    """Inserta la instruccion, la leyenda y la tabla de cada bloque declarado."""
    declaracion = m15.leer_declaracion_tablas(
        _RAIZ / "config" / "informe_tablas.yaml")
    hechos: list[str] = []

    for bloque in BLOQUES_NUEVOS:
        leyenda = str(bloque["leyenda"])
        entrada = declaracion.get(m15._normalizar_leyenda(leyenda))
        if entrada is None:
            hechos.append(f"NO ESTA DECLARADA en informe_tablas.yaml: "
                          f"{leyenda}")
            continue
        titulos = [str(t) for t in entrada.get("titulos") or []]
        columnas = [str(c) for c in entrada.get("columnas") or []]
        if len(titulos) != len(columnas):
            hechos.append(f"{leyenda}: {len(columnas)} columna(s) declarada(s) "
                          f"y {len(titulos)} titulo(s)")
            continue

        objetos = _indexar(documento)
        if _leyenda(objetos, leyenda)[0] >= 0:
            hechos.append(f"YA INSERTADA: {leyenda}")
            continue

        # --- de donde se copia el formato ---
        posicion, modelo_leyenda = _leyenda(
            objetos, "Datos tránsito crecientes")
        if modelo_leyenda is None or objetos[posicion + 1][0] != "t":
            hechos.append("NO SE ENCONTRO la tabla modelo de la que copiar el "
                          "formato ('Datos tránsito crecientes')")
            continue
        modelo_tabla = objetos[posicion + 1][1]
        modelo_instruccion = next(
            (o for c, o in objetos
             if c == "p" and m15.clasificar(o.text)[0] == "tabla"), None)
        modelo_fuente = next((o for c, o in objetos
                              if c == "p" and o.style.name == "Fuente"), None)
        if modelo_instruccion is None or modelo_fuente is None:
            hechos.append("NO SE ENCONTRO instruccion o linea de fuente de la "
                          "que copiar el formato")
            continue

        # --- donde va ---
        if bloque.get("tras_tabla"):
            indice, _ = _leyenda(objetos, str(bloque["tras_tabla"]))
            if indice < 0:
                hechos.append(f"NO SE ENCONTRO la tabla "
                              f"'{bloque['tras_tabla']}' tras la que insertar")
                continue
            # Detras de esa tabla y de su linea de fuente, no pegado a la
            # leyenda: entre la instruccion y su tabla no puede entrar nada, o
            # el M15 deja de encontrarla y la reporta 'sin tabla detras'.
            final = indice + 1
            while (final + 1 < len(objetos)
                   and not (objetos[final + 1][0] == "p"
                            and m15.clasificar(objetos[final + 1][1].text)[0])
                   and objetos[final + 1][0] != "t"
                   and objetos[final + 1][1].style.name in ("Fuente",
                                                            "Caption")):
                final += 1
            ancla, despues = objetos[final][1], True
        else:
            indice, encontrada = _leyenda(objetos, str(bloque["antes_de"]))
            if encontrada is None:
                hechos.append(f"NO SE ENCONTRO la leyenda "
                              f"'{bloque['antes_de']}' antes de la que "
                              f"insertar")
                continue
            ancla, despues = encontrada, False

        if not escribir:
            hechos.append(f"insertaria '{leyenda}' con {len(columnas)} "
                          f"columna(s) ({bloque['motivo']})")
            continue

        instruccion = _clonar(
            modelo_instruccion, INSTRUCCION.format(numero=bloque["numero"]))
        leyenda_nueva = _clonar_leyenda(modelo_leyenda, leyenda)
        tabla = _clonar_tabla(modelo_tabla, len(columnas))
        for celda, titulo in zip(tabla.rows[0].cells, titulos):
            m15._fijar_texto(celda, titulo)
        # LAS FILAS DE DATOS SE DEJAN EN BLANCO. La tabla modelo las trae con
        # los numeros del estudio del que se copio la plantilla, y el M15
        # escribe tantas filas como traiga la fuente: una que sobrara se
        # quedaria mezclando dos estudios sin que nada lo advirtiera.
        for fila in tabla.rows[1:]:
            for celda in fila.cells:
                m15._fijar_texto(celda, "")
        fuente = _clonar(modelo_fuente, str(bloque["fuente"]))

        elemento = ancla._element
        piezas = [instruccion, leyenda_nueva, tabla._tbl, fuente]
        if despues:
            # addnext deja cada pieza justo detras del ancla, de modo que hay
            # que ir moviendo el ancla o el bloque sale al reves.
            for pieza in piezas:
                elemento.addnext(pieza)
                elemento = pieza
        else:
            # addprevious deja cada pieza justo delante del ancla, que NO se
            # mueve: insertando en orden, cada una queda detras de la anterior.
            # Recorrerlas al reves invertia el bloque entero, y el resultado
            # era 'Fuente', tabla, leyenda, instruccion.
            for pieza in piezas:
                elemento.addprevious(pieza)
        hechos.append(f"insertada '{leyenda}' con {len(columnas)} columna(s) "
                      f"({bloque['motivo']})")
    return hechos


def main() -> int:
    analizador = argparse.ArgumentParser(description=__doc__)
    analizador.add_argument("--escribir", action="store_true",
                            help="aplica los cambios; sin esto solo los lista")
    argumentos = analizador.parse_args()

    hechos = consolidar(PLANTILLA, argumentos.escribir)
    verbo = "aplicado" if argumentos.escribir else "pendiente"
    for linea in hechos:
        print(f"  [{verbo}] {linea}")
    print(f"\n{len(hechos)} cambio(s).")
    if not argumentos.escribir:
        print("Nada se escribio. Repetir con --escribir.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
